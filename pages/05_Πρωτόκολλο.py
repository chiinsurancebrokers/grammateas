# -*- coding: utf-8 -*-
"""Σελίδα 05 — Πρωτόκολλο Εγγράφων (Άρθρο 37)"""
import sys; sys.path.append("..")
import io, base64, re, json
import streamlit as st
from datetime import date
from modules.database import init_db, get_protokollon, save_proto, next_proto_number, get_conn

init_db()

# ══════════════════════════════════════════════════════════════
# MIGRATION — προσθήκη columns αρχείου αν δεν υπάρχουν
# ══════════════════════════════════════════════════════════════
def _migrate():
    conn = get_conn()
    for col in ["αρχείο_bytes BLOB", "αρχείο_όνομα TEXT", "αρχείο_τύπος TEXT"]:
        try:
            conn.execute(f"ALTER TABLE πρωτόκολλο ADD COLUMN {col}")
        except Exception:
            pass
    conn.commit()
    conn.close()

_migrate()

# ══════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════
def get_anthropic_key():
    try:
        return (st.secrets.get("AI", {}).get("ANTHROPIC_API_KEY")
                or st.secrets.get("ANTHROPIC_API_KEY", ""))
    except Exception:
        return ""


def extract_text(file_bytes: bytes, filename: str, mime: str) -> str:
    """Εξαγωγή κειμένου από PDF ή DOCX."""
    try:
        if (mime or "").startswith("application/pdf") or filename.lower().endswith(".pdf"):
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n".join(page.get_text() for page in doc)[:4000]
        elif "wordprocessingml" in (mime or "") or filename.lower().endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    seen = set()
                    for cell in row.cells:
                        cid = id(cell._tc)
                        if cid not in seen and cell.text.strip():
                            lines.append(cell.text.strip())
                            seen.add(cid)
            return "\n".join(lines)[:4000]
    except Exception as e:
        return f"[Σφάλμα εξαγωγής: {e}]"
    return ""


def ai_extract(text: str, filename: str) -> dict:
    """Claude εξάγει μεταδεδομένα για αυτόματη συμπλήρωση."""
    key = get_anthropic_key()
    if not key or not text.strip():
        return {}
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=key)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=600,
            system=(
                "Εξάγεις μεταδεδομένα εγγράφου για Πρωτόκολλο Τεκτονικής Στοάς.\n"
                "Επέστρεψε ΜΟΝΟ έγκυρο JSON χωρίς markdown:\n"
                '{"θέμα":"...","ημερομηνία":"YYYY-MM-DD ή κενό",'
                '"αποστολέας":"...","παραλήπτης":"...",'
                '"κατεύθυνση":"Εισερχόμενο ή Εξερχόμενο","περίληψη":"..."}'
            ),
            messages=[{"role": "user", "content":
                f"Αρχείο: {filename}\n\nΚείμενο:\n{text[:2500]}"}],
        )
        raw = msg.content[0].text.strip() if msg.content else ""
        if not raw:
            return {}
        # Καθαρισμός markdown fences αν υπάρχουν
        raw = re.sub(r"^```[a-z]*\s*", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\s*```\s*$", "", raw, flags=re.MULTILINE).strip()
        return json.loads(raw)
    except Exception as e:
        st.warning(f"⚠️ AI extraction: {e}")
        return {}


def save_with_file(data: dict, fbytes=None, fname=None, ftype=None) -> int:
    if fbytes:
        data["αρχείο_bytes"] = fbytes
        data["αρχείο_όνομα"] = fname or ""
        data["αρχείο_τύπος"] = ftype or ""
    return save_proto(data)


def delete_proto(proto_id: int):
    conn = get_conn()
    conn.execute("DELETE FROM πρωτόκολλο WHERE id=?", (proto_id,))
    conn.commit()
    conn.close()


def get_proto_row(proto_id: int) -> dict:
    conn = get_conn()
    cur = conn.execute("SELECT * FROM πρωτόκολλο WHERE id=?", (proto_id,))
    cols = [d[0] for d in cur.description]
    row  = cur.fetchone()
    conn.close()
    return dict(zip(cols, row)) if row else {}


def get_file(proto_id: int):
    conn = get_conn()
    row = conn.execute(
        "SELECT αρχείο_bytes, αρχείο_όνομα, αρχείο_τύπος FROM πρωτόκολλο WHERE id=?",
        (proto_id,)
    ).fetchone()
    conn.close()
    return (row[0], row[1], row[2]) if row and row[0] else (None, None, None)


_preview_counter = [0]

def render_preview(file_bytes: bytes, filename: str, mime: str):
    """
    Προεπισκόπηση αρχείου.
    PDF → εξαγωγή κειμένου (το Chrome μπλοκάρει data: iframes).
    DOCX → εξαγωγή κειμένου.
    """
    _preview_counter[0] += 1
    text = extract_text(file_bytes, filename, mime or "")
    if text and not text.startswith("[Σφάλμα"):
        st.text_area("📄 Περιεχόμενο εγγράφου", value=text, height=420,
                     disabled=True, key=f"preview_text_{_preview_counter[0]}")
    else:
        st.info("Δεν ήταν δυνατή η προεπισκόπηση. Χρησιμοποιήστε το κουμπί λήψης.")


# ══════════════════════════════════════════════════════════════
# PAGE
# ══════════════════════════════════════════════════════════════
st.set_page_config(page_title="Πρωτόκολλο", page_icon="📬", layout="wide")
st.markdown("# 📬 Πρωτόκολλο Εγγράφων")
st.caption("Άρθρο 37 — Πρωτόκολλο εισερχομένων & εξερχομένων εγγράφων · Διεξαγωγή αλληλογραφίας")

tab_list, tab_new, tab_upload, tab_import = st.tabs([
    "📋 Πρωτόκολλο", "➕ Νέα Εγγραφή", "📎 Ανέβασμα Εγγράφου", "📥 Εισαγωγή CSV"
])

# ══════════════════════════════════════════════════════════════
# TAB 1 — ΛΙΣΤΑ
# ══════════════════════════════════════════════════════════════
with tab_list:
    c1, c2, c3, c4 = st.columns(4)
    with c1: year = st.selectbox("Έτος", list(range(date.today().year, 2019, -1)))
    with c2: direction = st.selectbox("Κατεύθυνση", ["Όλα", "Εισερχόμενο", "Εξερχόμενο"])
    with c3: status_f = st.selectbox("Κατάσταση", ["Όλες", "Εκκρεμές", "Απαντήθηκε", "Αρχειοθετήθηκε"])
    with c4: search = st.text_input("🔍 Θέμα")

    df = get_protokollon(year=year, direction="all" if direction == "Όλα" else direction)
    if status_f != "Όλες":
        df = df[df["κατάσταση"] == status_f]
    if search:
        df = df[df["θέμα"].str.contains(search, case=False, na=False)]

    ekk = len(df[df["κατάσταση"] == "Εκκρεμές"])
    if ekk:
        st.warning(f"⚠️ {ekk} εκκρεμή έγγραφα")

    if df.empty:
        st.info("Δεν υπάρχουν εγγραφές.")
    else:
        show_cols = ["αρ_πρωτ", "ημερομηνία", "κατεύθυνση",
                     "αποστολέας", "παραλήπτης", "θέμα", "κατάσταση"]
        if "αρχείο_όνομα" in df.columns:
            df["📎"] = df["αρχείο_όνομα"].apply(lambda x: "📎" if x and str(x).strip() else "")
            show_cols = ["📎"] + show_cols
        st.dataframe(df[show_cols], use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader("📄 Λήψη & Προεπισκόπηση Αρχείου")
    col_a, col_b = st.columns([1, 3])
    with col_a:
        view_id = st.number_input("ID Εγγραφής", min_value=1, step=1, key="view_id")
        if st.button("🔍 Φόρτωση", use_container_width=True):
            fb, fn, ft = get_file(int(view_id))
            if fb:
                st.session_state["pv_bytes"] = fb
                st.session_state["pv_name"]  = fn
                st.session_state["pv_type"]  = ft
            else:
                st.warning("Δεν υπάρχει αρχείο για αυτή την εγγραφή.")
    with col_b:
        if "pv_bytes" in st.session_state:
            fb = st.session_state["pv_bytes"]
            fn = st.session_state["pv_name"]
            ft = st.session_state["pv_type"]
            st.download_button(f"⬇️ Λήψη: {fn}", data=fb, file_name=fn,
                               mime=ft or "application/octet-stream",
                               use_container_width=True)
            with st.expander("👁️ Προεπισκόπηση κειμένου"):
                render_preview(fb, fn, ft)

    st.markdown("---")

    # ── Edit / Delete panel ───────────────────────────────────
    st.subheader("✏️ Επεξεργασία Εγγραφής")
    edit_id = st.number_input("ID Εγγραφής για επεξεργασία / διαγραφή",
                               min_value=1, step=1, key="edit_id")

    col_load, col_del = st.columns([2, 1])
    with col_load:
        load_btn = st.button("📂 Φόρτωση για επεξεργασία", use_container_width=True)
    with col_del:
        del_btn = st.button("🗑️ Διαγραφή", type="secondary", use_container_width=True)

    if del_btn:
        delete_proto(int(edit_id))
        st.success(f"✅ Η εγγραφή {edit_id} διαγράφηκε.")
        st.rerun()

    if load_btn or st.session_state.get("editing_proto_id") == int(edit_id):
        row = get_proto_row(int(edit_id))
        if not row:
            st.error("Δεν βρέθηκε εγγραφή με αυτό το ID.")
        else:
            st.session_state["editing_proto_id"] = int(edit_id)
            st.markdown(f"**Επεξεργασία ΑΠ {row.get('αρ_πρωτ','')}**")

            with st.form("edit_proto_form"):
                ec1, ec2, ec3 = st.columns(3)
                with ec1:
                    e_αρ   = st.text_input("Αρ. Πρωτ.", value=row.get("αρ_πρωτ",""))
                    e_ημερ = st.text_input("Ημερομηνία (YYYY-MM-DD)", value=row.get("ημερομηνία",""))
                with ec2:
                    dir_opts = ["Εισερχόμενο","Εξερχόμενο"]
                    dir_idx  = dir_opts.index(row.get("κατεύθυνση","Εισερχόμενο"))                                if row.get("κατεύθυνση") in dir_opts else 0
                    e_κατ  = st.radio("Κατεύθυνση", dir_opts, index=dir_idx, horizontal=True)
                    st_opts = ["Εκκρεμές","Απαντήθηκε","Αρχειοθετήθηκε"]
                    st_idx  = st_opts.index(row.get("κατάσταση","Εκκρεμές"))                                if row.get("κατάσταση") in st_opts else 0
                    e_st   = st.selectbox("Κατάσταση", st_opts, index=st_idx)
                with ec3:
                    e_αποστ = st.text_input("Αποστολέας", value=row.get("αποστολέας","") or "")
                    e_παραλ = st.text_input("Παραλήπτης", value=row.get("παραλήπτης","") or "")

                e_θέμα   = st.text_input("Θέμα", value=row.get("θέμα","") or "")
                e_περιγρ = st.text_area("Περιγραφή", value=row.get("περιγραφή","") or "", height=80)
                e_σχετ   = st.text_input("Σχετικό", value=row.get("αρ_σχετικού","") or "")
                e_notes  = st.text_area("Παρατηρήσεις", value=row.get("παρατηρήσεις","") or "", height=60)

                if st.form_submit_button("💾 Αποθήκευση αλλαγών", type="primary",
                                          use_container_width=True):
                    save_proto({
                        "id":           int(edit_id),
                        "αρ_πρωτ":      e_αρ,
                        "ημερομηνία":   e_ημερ,
                        "κατεύθυνση":   e_κατ,
                        "αποστολέας":   e_αποστ,
                        "παραλήπτης":   e_παραλ,
                        "θέμα":         e_θέμα,
                        "περιγραφή":    e_περιγρ,
                        "αρ_σχετικού":  e_σχετ,
                        "κατάσταση":    e_st,
                        "παρατηρήσεις": e_notes,
                        "ημ_απάντησης": str(date.today()) if e_st == "Απαντήθηκε" else None,
                    })
                    st.success(f"✅ Η εγγραφή {edit_id} ενημερώθηκε!")
                    st.session_state.pop("editing_proto_id", None)
                    st.rerun()

# ══════════════════════════════════════════════════════════════
# TAB 2 — ΝΕΑ ΕΓΓΡΑΦΗ (manual)
# ══════════════════════════════════════════════════════════════
with tab_new:
    st.subheader("➕ Νέα Εγγραφή Πρωτοκόλλου")
    auto_num = next_proto_number(date.today().year)

    with st.form("new_proto_form"):
        c1, c2, c3 = st.columns(3)
        with c1:
            αρ   = st.text_input("Αριθμός Πρωτοκόλλου", value=auto_num)
            ημερ = st.date_input("Ημερομηνία *", value=date.today())
        with c2:
            κατ    = st.radio("Κατεύθυνση *", ["Εισερχόμενο", "Εξερχόμενο"], horizontal=True)
            st_val = st.selectbox("Κατάσταση", ["Εκκρεμές", "Απαντήθηκε", "Αρχειοθετήθηκε"])
        with c3:
            αποστ = st.text_input("Αποστολέας")
            παραλ = st.text_input("Παραλήπτης")

        θέμα   = st.text_input("Θέμα *")
        περιγρ = st.text_area("Περιγραφή / Σύνοψη", height=80)
        σχετ   = st.text_input("Σχετικό (αρ. πρωτ.)")
        notes  = st.text_area("Παρατηρήσεις", height=60)
        attach = st.file_uploader("📎 Επισύναψη αρχείου (προαιρετικό)",
                                  type=["pdf", "docx"], key="manual_attach")

        if st.form_submit_button("💾 Καταχώρηση", use_container_width=True, type="primary"):
            if not θέμα:
                st.error("Το θέμα είναι υποχρεωτικό!")
            else:
                save_with_file(
                    {"αρ_πρωτ": αρ, "ημερομηνία": str(ημερ), "κατεύθυνση": κατ,
                     "αποστολέας": αποστ, "παραλήπτης": παραλ, "θέμα": θέμα,
                     "περιγραφή": περιγρ, "αρ_σχετικού": σχετ,
                     "κατάσταση": st_val, "παρατηρήσεις": notes},
                    fbytes=attach.read() if attach else None,
                    fname=attach.name if attach else None,
                    ftype=attach.type if attach else None,
                )
                st.success(f"✅ Εγγράφηκε με ΑΠ {αρ}!")
                st.rerun()

# ══════════════════════════════════════════════════════════════
# TAB 3 — ΑΝΕΒΑΣΜΑ + AI SCRAPING
# ══════════════════════════════════════════════════════════════
with tab_upload:
    st.subheader("📎 Ανέβασμα Εγγράφου & Αυτόματη Πρωτοκόλληση")
    st.caption(
        "Ανεβάστε PDF ή DOCX. Το σύστημα εξάγει αυτόματα θέμα, ημερομηνία, "
        "αποστολέα & παραλήπτη με AI και αποδίδει αριθμό πρωτοκόλλου."
    )

    uploaded = st.file_uploader("Επιλέξτε αρχείο (PDF ή DOCX)",
                                 type=["pdf", "docx"], key="upload_doc")

    if uploaded:
        file_bytes = uploaded.read()
        filename   = uploaded.name
        mime       = uploaded.type

        col_prev, col_form = st.columns([1, 1])

        # ── Προεπισκόπηση ─────────────────────────────────────
        with col_prev:
            st.markdown("#### 👁️ Προεπισκόπηση")
            render_preview(file_bytes, filename, mime)

        # ── Φόρμα ─────────────────────────────────────────────
        with col_form:
            st.markdown("#### 📋 Στοιχεία Πρωτοκόλλου")

            # AI extraction — τρέχει μία φορά ανά αρχείο
            cache_key = f"ai_{filename}_{len(file_bytes)}"
            if st.session_state.get("ai_cache_key") != cache_key:
                with st.spinner("🤖 AI εξαγωγή μεταδεδομένων…"):
                    doc_text = extract_text(file_bytes, filename, mime)
                    meta     = ai_extract(doc_text, filename)
                    st.session_state["ai_meta"]      = meta
                    st.session_state["ai_text"]      = doc_text
                    st.session_state["ai_cache_key"] = cache_key
            else:
                meta     = st.session_state.get("ai_meta", {})
                doc_text = st.session_state.get("ai_text", "")

            if meta:
                st.success("✅ AI εξαγωγή — ελέγξτε & διορθώστε:")

            # Αυτόματος αριθμός ΑΠ
            auto_num = next_proto_number(date.today().year)

            # Κατεύθυνση από AI
            dir_idx = 1 if meta.get("κατεύθυνση", "") == "Εξερχόμενο" else 0

            # Ημερομηνία από AI
            ai_date = date.today()
            try:
                from datetime import datetime as _dt
                if meta.get("ημερομηνία"):
                    ai_date = _dt.strptime(meta["ημερομηνία"], "%Y-%m-%d").date()
            except Exception:
                pass

            with st.form("upload_proto_form"):
                u_αρ    = st.text_input("Αριθμός Πρωτοκόλλου", value=auto_num)
                u_κατ   = st.radio("Κατεύθυνση *", ["Εισερχόμενο", "Εξερχόμενο"],
                                    index=dir_idx, horizontal=True)
                u_ημερ  = st.date_input("Ημερομηνία *", value=ai_date)
                u_αποστ = st.text_input("Αποστολέας",  value=meta.get("αποστολέας", ""))
                u_παραλ = st.text_input("Παραλήπτης",  value=meta.get("παραλήπτης", ""))
                u_θέμα  = st.text_input("Θέμα *",       value=meta.get("θέμα", filename))
                u_περιγρ= st.text_area("Περιγραφή / Σύνοψη",
                                        value=meta.get("περίληψη", ""), height=100)
                u_σχετ  = st.text_input("Σχετικό (αρ. πρωτ.)")
                u_st    = st.selectbox("Κατάσταση",
                                        ["Εκκρεμές", "Απαντήθηκε", "Αρχειοθετήθηκε"])
                u_notes = st.text_area("Παρατηρήσεις", height=60)

                if st.form_submit_button("💾 Καταχώρηση & Αρχειοθέτηση",
                                          use_container_width=True, type="primary"):
                    if not u_θέμα:
                        st.error("Το θέμα είναι υποχρεωτικό!")
                    else:
                        new_id = save_with_file(
                            {"αρ_πρωτ": u_αρ, "ημερομηνία": str(u_ημερ),
                             "κατεύθυνση": u_κατ, "αποστολέας": u_αποστ,
                             "παραλήπτης": u_παραλ, "θέμα": u_θέμα,
                             "περιγραφή": u_περιγρ, "αρ_σχετικού": u_σχετ,
                             "κατάσταση": u_st, "παρατηρήσεις": u_notes},
                            fbytes=file_bytes, fname=filename, ftype=mime,
                        )
                        # Καθαρισμός cache
                        for k in ["ai_meta", "ai_text", "ai_cache_key"]:
                            st.session_state.pop(k, None)
                        st.success(f"✅ **{filename}** → ΑΠ **{u_αρ}** (ID: {new_id})")
                        st.balloons()
                        st.rerun()

        # Raw κείμενο AI
        if st.session_state.get("ai_text"):
            with st.expander("🔍 Κείμενο που ανέγνωσε το AI"):
                st.text(st.session_state["ai_text"][:2000])


# ══════════════════════════════════════════════════════════════
# TAB 4 — ΕΙΣΑΓΩΓΗ CSV
# ══════════════════════════════════════════════════════════════
with tab_import:
    st.subheader("📥 Εισαγωγή από CSV")
    st.caption(
        "Ανεβάστε το CSV που εξάγατε από το πρωτόκολλο. "
        "Οι εγγραφές εισάγονται μόνο αν δεν υπάρχει ήδη ο ίδιος αρ. πρωτοκόλλου."
    )

    csv_file = st.file_uploader("Επιλέξτε CSV αρχείο", type=["csv"], key="import_csv")

    if csv_file:
        import pandas as pd
        try:
            df_csv = pd.read_csv(csv_file, encoding="utf-8-sig")
            # Καθαρισμός στήλης 📎 αν υπάρχει
            if "📎" in df_csv.columns:
                df_csv = df_csv.drop(columns=["📎"])

            st.markdown(f"**{len(df_csv)} εγγραφές βρέθηκαν στο CSV**")
            st.dataframe(df_csv.head(10), use_container_width=True, hide_index=True)

            col_imp, col_skip = st.columns(2)
            with col_imp:
                overwrite = st.checkbox("Αντικατάσταση αν υπάρχει ήδη ο ΑΠ", value=False)
            with col_skip:
                st.info("Χωρίς αντικατάσταση: παραλείπονται διπλότυπα")

            if st.button("📥 Εισαγωγή εγγραφών", type="primary", use_container_width=True):
                conn = get_conn()
                inserted = 0
                skipped  = 0
                updated  = 0

                for _, row in df_csv.iterrows():
                    αρ = str(row.get("αρ_πρωτ", "")).strip()
                    if not αρ:
                        skipped += 1
                        continue

                    # Έλεγχος αν υπάρχει ήδη
                    existing = conn.execute(
                        "SELECT id FROM πρωτόκολλο WHERE αρ_πρωτ=?", (αρ,)
                    ).fetchone()

                    record = {
                        "αρ_πρωτ":     αρ,
                        "ημερομηνία":  str(row.get("ημερομηνία", "")).strip(),
                        "κατεύθυνση":  str(row.get("κατεύθυνση", "Εισερχόμενο")).strip(),
                        "αποστολέας":  str(row.get("αποστολέας", "") or "").strip(),
                        "παραλήπτης":  str(row.get("παραλήπτης", "") or "").strip(),
                        "θέμα":        str(row.get("θέμα", "") or "").strip() or "—",
                        "περιγραφή":   str(row.get("περιγραφή", "") or "").strip(),
                        "αρ_σχετικού": str(row.get("αρ_σχετικού", "") or "").strip(),
                        "κατάσταση":   str(row.get("κατάσταση", "Εκκρεμές")).strip(),
                        "παρατηρήσεις":str(row.get("παρατηρήσεις", "") or "").strip(),
                    }

                    if existing:
                        if overwrite:
                            fields = ", ".join(f"{k}=?" for k in record)
                            conn.execute(
                                f"UPDATE πρωτόκολλο SET {fields} WHERE id=?",
                                (*record.values(), existing[0])
                            )
                            updated += 1
                        else:
                            skipped += 1
                    else:
                        cols = ", ".join(record.keys())
                        phs  = ", ".join("?" * len(record))
                        conn.execute(
                            f"INSERT INTO πρωτόκολλο ({cols}) VALUES ({phs})",
                            list(record.values())
                        )
                        inserted += 1

                conn.commit()
                conn.close()

                st.success(
                    f"✅ Ολοκληρώθηκε — "
                    f"**{inserted}** νέες | "
                    f"**{updated}** ενημερώθηκαν | "
                    f"**{skipped}** παραλείφθηκαν"
                )
                st.rerun()

        except Exception as e:
            st.error(f"❌ Σφάλμα ανάγνωσης CSV: {e}")

    st.markdown("---")
    st.subheader("📤 Εξαγωγή σε CSV")
    st.caption("Κατεβάστε όλο το πρωτόκολλο σε CSV για backup ή επαναφόρτωση.")

    year_exp = st.selectbox("Έτος", ["Όλα"] + list(range(date.today().year, 2019, -1)),
                             key="export_year")
    if st.button("📤 Εξαγωγή", use_container_width=True):
        import pandas as pd
        df_exp = get_protokollon(
            year=None if year_exp == "Όλα" else int(year_exp)
        )
        # Remove BLOB column for CSV
        for col in ["αρχείο_bytes", "αρχείο_όνομα", "αρχείο_τύπος"]:
            if col in df_exp.columns:
                df_exp = df_exp.drop(columns=[col])
        csv_data = df_exp.to_csv(index=False, encoding="utf-8-sig")
        st.download_button(
            "⬇️ Λήψη CSV",
            data=csv_data.encode("utf-8-sig"),
            file_name=f"πρωτόκολλο_{year_exp}_{date.today()}.csv",
            mime="text/csv",
            use_container_width=True,
        )
