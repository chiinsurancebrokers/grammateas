# -*- coding: utf-8 -*-
"""
Page 20 — Hybrid AI Transcription & Formal Minutes Generator

Final version for Streamlit.

Features:
  • Manual API key input in sidebar if Streamlit secrets fail
  • Reads OpenAI / Claude keys from:
      1. manual sidebar input
      2. Streamlit root secrets
      3. Streamlit [AI] nested secrets
      4. environment variables
  • Audio upload: m4a/mp3/wav/mp4/webm/mpeg/mpga
  • Automatic chunking for large audio using imageio-ffmpeg / ffmpeg
  • OpenAI transcription for each chunk
  • Re-composition into one full transcript
  • Structured JSON extraction with OpenAI
  • Claude auto-detect for formal minutes synthesis
  • OpenAI fallback if Claude is missing or fails
  • Word / PDF / TXT export

Recommended requirements.txt:
streamlit>=1.30
pandas>=2.0
plotly>=5.0
anthropic>=0.25
openai>=1.50.0
Pillow
reportlab>=4.0
python-docx>=1.0.0
imageio-ffmpeg>=0.4.9
"""

import html
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st
from openai import BadRequestError, OpenAI
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet


# ══════════════════════════════════════════════════════════════
# PAGE SETUP
# ══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Μεταγραφή & Πρακτικά AI",
    page_icon="🎤",
    layout="wide",
)

st.markdown("# 🎤 Μεταγραφή & Σύνθεση Πρακτικών με AI")
st.caption(
    "OpenAI transcription με αυτόματο chunking · Claude/OpenAI για σύνθεση επίσημου πρακτικού · Word/PDF/TXT export"
)


# ══════════════════════════════════════════════════════════════
# CONSTANTS
# ══════════════════════════════════════════════════════════════
SAFE_CHUNK_MB = 22
SUPPORTED_AUDIO_TYPES = ["m4a", "mp3", "wav", "mp4", "webm", "mpeg", "mpga"]
DEFAULT_STOAA = "Σ∴ Στ∴ ΑΚΡΟΠΟΛΙΣ υπ’ αρ. 84"


# ══════════════════════════════════════════════════════════════
# API KEY HELPERS
# ══════════════════════════════════════════════════════════════
def get_secret(name: str) -> str:
    """
    Reads API key in this order:
    1. Manual sidebar input saved in session_state
    2. Streamlit root secrets
    3. Streamlit nested [AI] secrets
    4. Environment variables
    """
    manual_map = {
        "OPENAI_API_KEY": "MANUAL_OPENAI_API_KEY",
        "ANTHROPIC_API_KEY": "MANUAL_ANTHROPIC_API_KEY",
    }

    manual_key = manual_map.get(name)
    if manual_key and st.session_state.get(manual_key):
        return str(st.session_state[manual_key]).strip()

    try:
        root_value = st.secrets.get(name, "")
        if root_value:
            return str(root_value).strip()
    except Exception:
        pass

    try:
        ai_section = st.secrets.get("AI", {})
        if isinstance(ai_section, dict):
            nested_value = ai_section.get(name, "")
            if nested_value:
                return str(nested_value).strip()
    except Exception:
        pass

    return os.getenv(name, "").strip()


def mask_key(key: str) -> str:
    if not key:
        return "not set"
    if len(key) <= 12:
        return "set"
    return f"{key[:7]}...{key[-4:]}"


def get_openai_client() -> OpenAI:
    key = get_secret("OPENAI_API_KEY")
    if not key:
        st.error("❌ Δεν βρέθηκε OPENAI_API_KEY. Βάλ’ το στο sidebar ή στα Streamlit secrets.")
        st.stop()
    return OpenAI(api_key=key)


def get_anthropic_client():
    key = get_secret("ANTHROPIC_API_KEY")
    if not key:
        return None

    try:
        import anthropic
        return anthropic.Anthropic(api_key=key)
    except Exception as exc:
        st.warning(f"⚠️ Βρέθηκε Claude API key, αλλά δεν φορτώθηκε το anthropic package: {exc}")
        return None


def show_api_status() -> None:
    openai_key = get_secret("OPENAI_API_KEY")
    anthropic_key = get_secret("ANTHROPIC_API_KEY")

    st.markdown("### 🔎 API Status")

    if openai_key:
        st.success(f"✅ OpenAI detected: {mask_key(openai_key)}")
    else:
        st.error("❌ OpenAI API missing")

    if anthropic_key:
        st.success(f"✅ Claude detected: {mask_key(anthropic_key)}")
    else:
        st.info("ℹ️ Claude missing — θα γίνει fallback σε OpenAI για τη σύνθεση πρακτικού")


# ══════════════════════════════════════════════════════════════
# FFMPEG HELPERS
# ══════════════════════════════════════════════════════════════
def get_ffmpeg_path() -> str:
    """Uses bundled ffmpeg from imageio-ffmpeg where available."""
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception:
        ffmpeg = shutil.which("ffmpeg")
        if ffmpeg:
            return ffmpeg

    st.error(
        "❌ Δεν βρέθηκε ffmpeg. Πρόσθεσε `imageio-ffmpeg>=0.4.9` στο requirements.txt και κάνε redeploy."
    )
    st.stop()


def run_ffmpeg(cmd: List[str]) -> subprocess.CompletedProcess:
    return subprocess.run(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        check=False,
    )


def get_audio_duration_seconds(input_path: str) -> float:
    ffmpeg = get_ffmpeg_path()
    proc = run_ffmpeg([ffmpeg, "-i", input_path])
    text = proc.stderr or ""

    match = re.search(r"Duration:\s*(\d+):(\d+):(\d+(?:\.\d+)?)", text)
    if not match:
        return 0.0

    hours = int(match.group(1))
    minutes = int(match.group(2))
    seconds = float(match.group(3))
    return hours * 3600 + minutes * 60 + seconds


def split_audio_to_chunks(input_path: str, output_dir: str, target_mb: int = SAFE_CHUNK_MB) -> List[str]:
    """Converts and splits audio to MP3 mono 16k chunks under safe upload size."""
    input_size_mb = os.path.getsize(input_path) / (1024 * 1024)

    if input_size_mb <= target_mb:
        return [input_path]

    ffmpeg = get_ffmpeg_path()
    duration = get_audio_duration_seconds(input_path)

    if duration <= 0:
        segment_time = 8 * 60
    else:
        mb_per_second = input_size_mb / duration
        estimated_seconds = int((target_mb * 0.82) / max(mb_per_second, 0.001))
        segment_time = max(180, min(estimated_seconds, 12 * 60))

    chunk_pattern = os.path.join(output_dir, "chunk_%03d.mp3")

    cmd = [
        ffmpeg,
        "-y",
        "-i", input_path,
        "-vn",
        "-ac", "1",
        "-ar", "16000",
        "-b:a", "48k",
        "-f", "segment",
        "-segment_time", str(segment_time),
        "-reset_timestamps", "1",
        chunk_pattern,
    ]

    proc = run_ffmpeg(cmd)
    if proc.returncode != 0:
        st.error("❌ Αποτυχία στο automatic chunking με ffmpeg.")
        with st.expander("Τεχνικές λεπτομέρειες ffmpeg"):
            st.code(proc.stderr[-4000:])
        st.stop()

    chunks = sorted(str(p) for p in Path(output_dir).glob("chunk_*.mp3"))

    too_large = [p for p in chunks if os.path.getsize(p) / (1024 * 1024) > target_mb]
    if too_large:
        for p in chunks:
            try:
                os.remove(p)
            except Exception:
                pass

        smaller_time = max(120, segment_time // 2)
        chunk_pattern = os.path.join(output_dir, "chunk_small_%03d.mp3")

        cmd = [
            ffmpeg,
            "-y",
            "-i", input_path,
            "-vn",
            "-ac", "1",
            "-ar", "16000",
            "-b:a", "48k",
            "-f", "segment",
            "-segment_time", str(smaller_time),
            "-reset_timestamps", "1",
            chunk_pattern,
        ]

        proc = run_ffmpeg(cmd)
        if proc.returncode != 0:
            st.error("❌ Αποτυχία στο δεύτερο chunking attempt.")
            with st.expander("Τεχνικές λεπτομέρειες ffmpeg"):
                st.code(proc.stderr[-4000:])
            st.stop()

        chunks = sorted(str(p) for p in Path(output_dir).glob("chunk_small_*.mp3"))

    if not chunks:
        st.error("❌ Δεν δημιουργήθηκαν audio chunks.")
        st.stop()

    return chunks


# ══════════════════════════════════════════════════════════════
# TRANSCRIPTION
# ══════════════════════════════════════════════════════════════
def transcribe_single_file(client: OpenAI, path: str, language: str = "el") -> str:
    """Transcribes one audio file/chunk using OpenAI."""
    size_mb = os.path.getsize(path) / (1024 * 1024)

    if size_mb > 24.5:
        st.error(f"❌ Το chunk είναι ακόμη πολύ μεγάλο ({size_mb:.2f}MB).")
        st.stop()

    try:
        with open(path, "rb") as audio:
            result = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio,
                language=language,
                response_format="text",
            )
        return str(result or "").strip()

    except BadRequestError as exc:
        st.error("❌ Το OpenAI transcription API δεν δέχθηκε ένα audio chunk.")
        st.info("Συνήθεις αιτίες: χαλασμένο αρχείο, μη υποστηριζόμενο format ή chunk πάνω από 25MB.")
        with st.expander("Τεχνικές λεπτομέρειες"):
            st.code(str(exc))
        st.stop()

    except Exception as exc:
        st.error("❌ Αποτυχία κατά τη μεταγραφή.")
        with st.expander("Τεχνικές λεπτομέρειες"):
            st.code(str(exc))
        st.stop()


def transcribe_with_chunking(client: OpenAI, uploaded_file, language: str = "el") -> Tuple[str, List[Dict[str, Any]]]:
    """Save upload → split if needed → transcribe all chunks → recombine transcript."""
    original_suffix = Path(uploaded_file.name).suffix.lower() or ".m4a"

    with tempfile.TemporaryDirectory() as workdir:
        input_path = os.path.join(workdir, f"input{original_suffix}")
        with open(input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        chunk_dir = os.path.join(workdir, "chunks")
        os.makedirs(chunk_dir, exist_ok=True)

        chunks = split_audio_to_chunks(input_path, chunk_dir, target_mb=SAFE_CHUNK_MB)

        transcript_parts: List[str] = []
        chunk_log: List[Dict[str, Any]] = []

        progress = st.progress(0)
        status = st.empty()

        for idx, chunk_path in enumerate(chunks, start=1):
            size_mb = os.path.getsize(chunk_path) / (1024 * 1024)
            status.info(f"🎧 Μεταγραφή chunk {idx}/{len(chunks)} ({size_mb:.2f}MB)...")

            part = transcribe_single_file(client, chunk_path, language=language)
            transcript_parts.append(f"[CHUNK {idx}]\n{part}".strip())

            chunk_log.append({
                "chunk": idx,
                "file": os.path.basename(chunk_path),
                "size_mb": round(size_mb, 2),
                "characters": len(part),
            })

            progress.progress(int((idx / len(chunks)) * 100))

        status.success("✅ Η μεταγραφή όλων των chunks ολοκληρώθηκε.")
        return "\n\n".join(transcript_parts), chunk_log


# ══════════════════════════════════════════════════════════════
# STRUCTURED JSON EXTRACTION
# ══════════════════════════════════════════════════════════════
def safe_json_loads(raw: str) -> Dict[str, Any]:
    raw = (raw or "").strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw).strip()

    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            return data
    except Exception:
        pass

    return {
        "meeting": {},
        "speakers": [],
        "present": [],
        "absent": [],
        "dates": [],
        "agenda": [],
        "decisions": [],
        "actions": [],
        "financial_items": [],
        "uncertainties": ["Η AI ανάλυση δεν επέστρεψε valid JSON."],
        "summary": raw or "",
    }


def structured_extract_with_openai(client: OpenAI, transcript: str, meta: Dict[str, str]) -> Dict[str, Any]:
    system = """
Είσαι έμπειρος Γραμματεύς Τεκτονικής Στοάς και κάνεις δομημένη ανάλυση μεταγραφής.

Επέστρεψε ΜΟΝΟ valid JSON, χωρίς markdown, χωρίς σχόλια.

Ακριβές schema:
{
  "meeting": {
    "stoaa": "",
    "date": "",
    "degree": "",
    "chair": "",
    "secretary": ""
  },
  "speakers": [
    {"speaker": "Ομιλητής 1 ή όνομα αν προκύπτει", "points": ["..."]}
  ],
  "present": ["..."],
  "absent": ["..."],
  "dates": ["..."],
  "agenda": ["..."],
  "decisions": ["..."],
  "actions": [
    {"action": "...", "owner": "...", "deadline": "..."}
  ],
  "financial_items": ["..."],
  "uncertainties": ["..."],
  "summary": "Σύντομη σύνοψη"
}

Κανόνες:
- Μην επινοείς ονόματα ή αποφάσεις.
- Αν κάτι δεν προκύπτει, βάλε το στο "uncertainties".
- Αγνόησε τεχνικά markers όπως [CHUNK 1].
- Οι αποφάσεις να είναι καθαρές και πρακτικές.
- Οι ενέργειες να έχουν owner/deadline μόνο αν αναφέρονται.
""".strip()

    user_content = f"""
Στοιχεία από sidebar:
{json.dumps(meta, ensure_ascii=False, indent=2)}

Μεταγραφή:
{transcript}
""".strip()

    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.1,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user_content},
        ],
    )

    return safe_json_loads(response.choices[0].message.content)


# ══════════════════════════════════════════════════════════════
# MINUTES SYNTHESIS — CLAUDE FIRST, OPENAI FALLBACK
# ══════════════════════════════════════════════════════════════
def build_minutes_prompt(transcript: str, analysis: Dict[str, Any], meta: Dict[str, str]) -> str:
    return f"""
Συνέταξε επίσημο πρακτικό συνεδρίας στα Ελληνικά, με φυσικό ανθρώπινο ύφος και τεκτονική διατύπωση.

ΣΤΟΙΧΕΙΑ ΣΥΝΕΔΡΙΑΣ
{json.dumps(meta, ensure_ascii=False, indent=2)}

ΔΟΜΗΜΕΝΗ ΑΝΑΛΥΣΗ JSON
{json.dumps(analysis, ensure_ascii=False, indent=2)}

ΕΝΙΑΙΑ ΜΕΤΑΓΡΑΦΗ
{transcript}

ΑΠΑΙΤΟΥΜΕΝΗ ΔΟΜΗ ΠΡΑΚΤΙΚΟΥ:
ΠΡΑΚΤΙΚΟΝ ΣΥΝΕΔΡΙΑΣ
Σ∴ Στ∴ ...
Ημερομηνία — Βαθμός

ΕΝΑΡΞΙΣ
Να αναφέρεται η ημερομηνία, ο τόπος/Στοά, ο βαθμός εργασιών και, εφόσον προκύπτει, υπό ποιον τελέστηκε η συνεδρία.

ΠΑΡΟΝΤΕΣ
Να συμπεριληφθούν όσοι προκύπτουν σαφώς. Αν δεν προκύπτουν σαφώς, γράψε διακριτικά ότι εκ της μεταγραφής δεν προκύπτει πλήρης κατάλογος παρόντων.

ΘΕΜΑΤΑ ΗΜΕΡΗΣΙΑΣ ΔΙΑΤΑΞΕΩΣ
Να παρουσιαστούν τα βασικά θέματα.

ΣΥΖΗΤΗΣΙΣ
Να αποδοθεί καθαρά η ουσία, χωρίς περιττές επαναλήψεις ή προφορικότητες.

ΑΠΟΦΑΣΕΙΣ
Να γραφούν ξεκάθαρα, αριθμημένα όπου είναι χρήσιμο.

ΕΝΕΡΓΕΙΕΣ / ΕΚΚΡΕΜΟΤΗΤΕΣ
Να γραφούν πρακτικά, με υπεύθυνο και προθεσμία μόνο αν προκύπτουν.

ΛΗΞΙΣ
Σύντομη τυπική λήξη.

ΚΑΝΟΝΕΣ:
- Μην επινοείς γεγονότα.
- Μην γράψεις ότι το κείμενο δημιουργήθηκε από AI.
- Μην αναφέρεις markers όπως [CHUNK 1].
- Μην αφήνεις hallucinated ονόματα.
- Αν υπάρχει αβεβαιότητα, γράψε: "εκ της μεταγραφής δεν προκύπτει σαφώς".
- Χρησιμοποίησε τεκτονικές συντομογραφίες με μέτρο: Σ∴ Στ∴, Σεβ∴, Αδ∴, Διδ∴ όπου αρμόζει.
- Το αποτέλεσμα να είναι έτοιμο για επεξεργασία από Γραμματέα.
""".strip()


def draft_minutes_with_claude(anthropic_client, transcript: str, analysis: Dict[str, Any], meta: Dict[str, str]) -> Optional[str]:
    if anthropic_client is None:
        return None

    prompt = build_minutes_prompt(transcript, analysis, meta)

    model_candidates = [
        "claude-sonnet-4-5",
        "claude-3-5-sonnet-20241022",
    ]

    last_error = None
    for model_name in model_candidates:
        try:
            message = anthropic_client.messages.create(
                model=model_name,
                max_tokens=6000,
                temperature=0.2,
                system=(
                    "Είσαι έμπειρος Γραμματεύς Τεκτονικής Στοάς. "
                    "Συντάσσεις επίσημα πρακτικά σε φυσικό, ανθρώπινο, τεκτονικό ύφος."
                ),
                messages=[{"role": "user", "content": prompt}],
            )

            if message.content and hasattr(message.content[0], "text"):
                st.caption(f"Claude model used: {model_name}")
                return message.content[0].text.strip()

        except Exception as exc:
            last_error = exc
            continue

    if last_error:
        st.warning(f"⚠️ Claude failed, fallback σε OpenAI. Λεπτομέρεια: {last_error}")

    return None


def draft_minutes_with_openai(client: OpenAI, transcript: str, analysis: Dict[str, Any], meta: Dict[str, str]) -> str:
    prompt = build_minutes_prompt(transcript, analysis, meta)

    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.2,
        messages=[
            {
                "role": "system",
                "content": (
                    "Είσαι έμπειρος Γραμματεύς Τεκτονικής Στοάς. "
                    "Συντάσσεις επίσημα πρακτικά σε φυσικό, ανθρώπινο, τεκτονικό ύφος."
                ),
            },
            {"role": "user", "content": prompt},
        ],
    )

    return response.choices[0].message.content.strip()


def draft_minutes_hybrid(
    openai_client: OpenAI,
    anthropic_client,
    transcript: str,
    analysis: Dict[str, Any],
    meta: Dict[str, str],
) -> Tuple[str, str]:
    claude_text = draft_minutes_with_claude(anthropic_client, transcript, analysis, meta)
    if claude_text:
        return claude_text, "Claude"

    openai_text = draft_minutes_with_openai(openai_client, transcript, analysis, meta)
    return openai_text, "OpenAI fallback"


# ══════════════════════════════════════════════════════════════
# EXPORT HELPERS
# ══════════════════════════════════════════════════════════════
def create_docx(minutes_text: str, analysis: Dict[str, Any], meta: Dict[str, str]) -> bytes:
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    title = doc.add_heading("ΠΡΑΚΤΙΚΟΝ ΣΥΝΕΔΡΙΑΣ", level=1)
    title.alignment = 1

    doc.add_paragraph(f"Στοά: {meta.get('stoaa', '')}")
    doc.add_paragraph(f"Ημερομηνία: {meta.get('meeting_date', '')}")
    doc.add_paragraph(f"Βαθμός: {meta.get('degree', '')}")
    doc.add_paragraph("")

    for line in minutes_text.splitlines():
        clean = line.strip()
        if not clean:
            doc.add_paragraph("")
        elif clean.isupper() and len(clean) < 90:
            doc.add_heading(clean, level=2)
        else:
            doc.add_paragraph(clean)

    doc.add_page_break()
    doc.add_heading("Δομημένη AI Ανάλυση", level=1)

    sections = [
        ("Παρόντες", analysis.get("present", [])),
        ("Απόντες", analysis.get("absent", [])),
        ("Ημερομηνίες", analysis.get("dates", [])),
        ("Θέματα", analysis.get("agenda", [])),
        ("Αποφάσεις", analysis.get("decisions", [])),
        ("Οικονομικά", analysis.get("financial_items", [])),
        ("Αβεβαιότητες", analysis.get("uncertainties", [])),
    ]

    for heading, items in sections:
        doc.add_heading(heading, level=2)
        if items:
            for item in items:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph("Δεν προκύπτει σαφώς.")

    actions = analysis.get("actions", [])
    doc.add_heading("Ενέργειες / Εκκρεμότητες", level=2)
    if actions:
        for action in actions:
            if isinstance(action, dict):
                doc.add_paragraph(
                    f"{action.get('action', '')} | Υπεύθυνος: {action.get('owner', '—')} | "
                    f"Προθεσμία: {action.get('deadline', '—')}"
                )
            else:
                doc.add_paragraph(str(action))
    else:
        doc.add_paragraph("Δεν προκύπτει σαφώς.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def register_greek_pdf_font() -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ]

    for font_path in candidates:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("GreekFont", font_path))
                return "GreekFont"
            except Exception:
                pass

    return "Helvetica"


def create_pdf(minutes_text: str, meta: Dict[str, str]) -> bytes:
    font_name = register_greek_pdf_font()

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()

    normal = ParagraphStyle(
        "GreekNormal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )
    title = ParagraphStyle(
        "GreekTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=16,
        leading=20,
    )

    story = [
        Paragraph("ΠΡΑΚΤΙΚΟΝ ΣΥΝΕΔΡΙΑΣ", title),
        Spacer(1, 12),
        Paragraph(f"Στοά: {html.escape(meta.get('stoaa', ''))}", normal),
        Paragraph(f"Ημερομηνία: {html.escape(meta.get('meeting_date', ''))}", normal),
        Paragraph(f"Βαθμός: {html.escape(meta.get('degree', ''))}", normal),
        Spacer(1, 12),
    ]

    safe_text = html.escape(minutes_text).replace("\n", "<br/>")
    story.append(Paragraph(safe_text, normal))

    pdf.build(story)
    buf.seek(0)
    return buf.read()


def make_safe_filename(prefix: str, extension: str) -> str:
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{prefix}_{stamp}.{extension}"


# ══════════════════════════════════════════════════════════════
# DISPLAY HELPERS
# ══════════════════════════════════════════════════════════════
def display_analysis(analysis: Dict[str, Any]) -> None:
    st.markdown("## 🔎 Δομημένη AI Ανάλυση")

    meeting = analysis.get("meeting", {}) or {}
    if meeting:
        with st.expander("🏛️ Στοιχεία Συνεδρίας από AI"):
            st.json(meeting)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 👥 Παρόντες")
        for item in analysis.get("present", []) or ["Δεν προκύπτει σαφώς"]:
            st.markdown(f"- {item}")

        st.markdown("### 📅 Ημερομηνίες")
        for item in analysis.get("dates", []) or ["Δεν προκύπτει σαφώς"]:
            st.markdown(f"- {item}")

        st.markdown("### 📋 Θέματα")
        for item in analysis.get("agenda", []) or ["Δεν προκύπτει σαφώς"]:
            st.markdown(f"- {item}")

    with col2:
        st.markdown("### ✅ Αποφάσεις")
        for item in analysis.get("decisions", []) or ["Δεν προκύπτει σαφώς"]:
            st.markdown(f"- {item}")

        st.markdown("### 📌 Ενέργειες / Εκκρεμότητες")
        actions = analysis.get("actions", [])
        if actions:
            for action in actions:
                if isinstance(action, dict):
                    st.markdown(
                        f"- **{action.get('action', '')}** "
                        f"| Υπεύθυνος: {action.get('owner', '—')} "
                        f"| Προθεσμία: {action.get('deadline', '—')}"
                    )
                else:
                    st.markdown(f"- {action}")
        else:
            st.markdown("- Δεν προκύπτει σαφώς")

        st.markdown("### ⚠️ Αβεβαιότητες")
        for item in analysis.get("uncertainties", []) or ["Δεν υπάρχουν εμφανείς αβεβαιότητες."]:
            st.markdown(f"- {item}")

    with st.expander("🗣️ Διαχωρισμός ανά ομιλητή"):
        speakers = analysis.get("speakers", [])
        if speakers:
            for speaker in speakers:
                st.markdown(f"**{speaker.get('speaker', 'Ομιλητής')}**")
                for point in speaker.get("points", []):
                    st.markdown(f"- {point}")
        else:
            st.caption("Δεν προέκυψε καθαρός διαχωρισμός ομιλητών.")

    with st.expander("🧾 Raw JSON"):
        st.json(analysis)


# ══════════════════════════════════════════════════════════════
# SIDEBAR UI
# ══════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Ρυθμίσεις")
    stoaa = st.text_input("Στοά", value=DEFAULT_STOAA)
    meeting_date = st.date_input("Ημερομηνία συνεδρίας", value=date.today())
    degree = st.selectbox("Βαθμός", ["Μαθητού", "Εταίρου", "Διδασκάλου", "Άλλο"])
    location = st.text_input("Τόπος", value="Τεκτ∴ Ναός")
    chair = st.text_input("Προεδρεύων / Σεβάσμιος", value="")
    secretary = st.text_input("Γραμματεύς", value="")
    language = st.selectbox("Γλώσσα μεταγραφής", ["el", "en"], index=0)

    st.markdown("---")
    st.markdown("## 🔐 API Keys")
    st.caption("Αν τα Streamlit secrets δεν διαβάζονται, βάλε τα keys προσωρινά εδώ.")

    manual_openai_key = st.text_input(
        "OpenAI API Key",
        value=st.session_state.get("MANUAL_OPENAI_API_KEY", ""),
        type="password",
        help="Απαραίτητο για transcription και fallback ανάλυση/σύνθεση.",
    )

    manual_anthropic_key = st.text_input(
        "Claude / Anthropic API Key",
        value=st.session_state.get("MANUAL_ANTHROPIC_API_KEY", ""),
        type="password",
        help="Προαιρετικό. Αν υπάρχει, χρησιμοποιείται για καλύτερη σύνθεση πρακτικού.",
    )

    if manual_openai_key:
        st.session_state["MANUAL_OPENAI_API_KEY"] = manual_openai_key.strip()

    if manual_anthropic_key:
        st.session_state["MANUAL_ANTHROPIC_API_KEY"] = manual_anthropic_key.strip()

    if st.button("🧹 Clear manual API keys", use_container_width=True):
        st.session_state.pop("MANUAL_OPENAI_API_KEY", None)
        st.session_state.pop("MANUAL_ANTHROPIC_API_KEY", None)
        st.rerun()

    st.markdown("---")
    show_api_status()


meta = {
    "stoaa": stoaa,
    "meeting_date": meeting_date.strftime("%d/%m/%Y"),
    "degree": degree,
    "location": location,
    "chair": chair,
    "secretary": secretary,
}


# ══════════════════════════════════════════════════════════════
# MAIN UI
# ══════════════════════════════════════════════════════════════
st.markdown("## 1️⃣ Upload Ηχογράφησης")
uploaded = st.file_uploader(
    "Ανεβάστε ηχογράφηση",
    type=SUPPORTED_AUDIO_TYPES,
    help="Αν το αρχείο είναι μεγάλο, θα κοπεί αυτόματα σε chunks και θα ξαναενωθεί η μεταγραφή.",
)

if uploaded:
    size_mb = uploaded.size / (1024 * 1024)
    st.audio(uploaded)
    st.success(f"✅ Φορτώθηκε: {uploaded.name} ({size_mb:.2f}MB)")

    if size_mb > SAFE_CHUNK_MB:
        st.warning(
            f"⚠️ Το αρχείο είναι πάνω από {SAFE_CHUNK_MB}MB. "
            "Θα γίνει αυτόματο chunking με ffmpeg και μετά ενιαία επανασύνθεση της μεταγραφής."
        )

    if st.button("🎧 Βήμα 1 · Μεταγραφή με automatic chunking", type="primary", use_container_width=True):
        openai_client = get_openai_client()

        with st.spinner("Γίνεται chunking και μεταγραφή..."):
            transcript, chunk_log = transcribe_with_chunking(openai_client, uploaded, language=language)

        st.session_state["page20_transcript"] = transcript
        st.session_state["page20_chunk_log"] = chunk_log
        st.session_state.pop("page20_analysis", None)
        st.session_state.pop("page20_minutes", None)

        st.success("✅ Η ενιαία μεταγραφή ολοκληρώθηκε.")


if "page20_chunk_log" in st.session_state:
    with st.expander("🧩 Chunks που μεταγράφηκαν"):
        st.dataframe(st.session_state["page20_chunk_log"], use_container_width=True)


if "page20_transcript" in st.session_state:
    st.markdown("---")
    st.markdown("## 2️⃣ Ενιαία Μεταγραφή")

    transcript_edit = st.text_area(
        "Μπορείτε να διορθώσετε τη μεταγραφή πριν την ανάλυση.",
        value=st.session_state["page20_transcript"],
        height=360,
    )
    st.session_state["page20_transcript"] = transcript_edit

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            "⬇️ Λήψη μεταγραφής TXT",
            data=transcript_edit.encode("utf-8"),
            file_name=make_safe_filename("transcription", "txt"),
            mime="text/plain",
            use_container_width=True,
        )

    with col2:
        if st.button("🔎 Βήμα 2 · Structured JSON Analysis", use_container_width=True):
            openai_client = get_openai_client()

            with st.spinner("Γίνεται δομημένη ανάλυση σε JSON..."):
                analysis = structured_extract_with_openai(openai_client, transcript_edit, meta)

            st.session_state["page20_analysis"] = analysis
            st.session_state.pop("page20_minutes", None)
            st.success("✅ Η δομημένη ανάλυση ολοκληρώθηκε.")


if "page20_analysis" in st.session_state:
    st.markdown("---")
    analysis = st.session_state["page20_analysis"]
    display_analysis(analysis)

    if st.button("📜 Βήμα 3 · Σύνθεση Επίσημου Πρακτικού", type="primary", use_container_width=True):
        openai_client = get_openai_client()
        anthropic_client = get_anthropic_client()

        with st.spinner("Συντίθεται το επίσημο πρακτικό με Claude/OpenAI fallback..."):
            minutes, model_used = draft_minutes_hybrid(
                openai_client=openai_client,
                anthropic_client=anthropic_client,
                transcript=st.session_state["page20_transcript"],
                analysis=analysis,
                meta=meta,
            )

        st.session_state["page20_minutes"] = minutes
        st.session_state["page20_model_used"] = model_used
        st.success(f"✅ Το πρακτικό δημιουργήθηκε με: {model_used}")


if "page20_minutes" in st.session_state:
    st.markdown("---")
    st.markdown("## 4️⃣ Επίσημο Πρακτικό")

    model_used = st.session_state.get("page20_model_used", "AI")
    st.caption(f"Σύνθεση με: {model_used}")

    minutes_edit = st.text_area(
        "Τελική επεξεργασία πριν το export.",
        value=st.session_state["page20_minutes"],
        height=520,
    )
    st.session_state["page20_minutes"] = minutes_edit

    analysis = st.session_state.get("page20_analysis", {})
    docx_bytes = create_docx(minutes_edit, analysis, meta)
    pdf_bytes = create_pdf(minutes_edit, meta)

    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            "⬇️ Word (.docx)",
            data=docx_bytes,
            file_name=make_safe_filename("praktiko", "docx"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with col2:
        st.download_button(
            "⬇️ PDF",
            data=pdf_bytes,
            file_name=make_safe_filename("praktiko", "pdf"),
            mime="application/pdf",
            use_container_width=True,
        )

    with col3:
        st.download_button(
            "⬇️ TXT",
            data=minutes_edit.encode("utf-8"),
            file_name=make_safe_filename("praktiko", "txt"),
            mime="text/plain",
            use_container_width=True,
        )


st.markdown("---")
with st.expander("ℹ️ Σημειώσεις λειτουργίας"):
    st.markdown(
        """
- Η μεταγραφή γίνεται με **OpenAI**.
- Αν το Streamlit δεν διαβάζει secrets, βάλε το OpenAI key χειροκίνητα στο sidebar.
- Η τελική σύνθεση πρακτικού γίνεται με **Claude**, εφόσον υπάρχει `ANTHROPIC_API_KEY`.
- Αν Claude δεν υπάρχει ή αποτύχει, γίνεται fallback σε **OpenAI**.
- Για μεγάλα audio, το αρχείο κόβεται αυτόματα σε μικρότερα chunks και μετά η μεταγραφή ενώνεται σε ενιαίο κείμενο.
- Το Word περιλαμβάνει και ξεχωριστή ενότητα με τη δομημένη AI ανάλυση.
        """
    )
