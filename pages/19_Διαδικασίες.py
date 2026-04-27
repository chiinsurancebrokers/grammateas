# -*- coding: utf-8 -*-
"""
Σελίδα 19 — Διαδικασίες & Έντυπα ΜΣΤΕ

Βοηθός συμπλήρωσης εντύπων βάσει:
  • Οδηγού Διαδικασιών & Απαιτούμενων Ενεργειών (Δεκ. 2017)
  • Γενικού Κανονισμού & Καταστατικού Χάρτη ΜΣΤΕ
  • Δεδομένων Μητρώου Στοάς

Υποστηριζόμενες διαδικασίες:
  1. Εισδοχή Υποψηφίου
  2. Αύξηση Μισθοδοσίας (Εταίρου / Διδασκάλου)
  3. Εκλογές Αξιωματικών & Εγκατάσταση
  4. Αναγγελία Μεταβολών
  5. Έγκριση Κοινής Συνεδρίας
  6. Έγκριση Λευκής Εορτής / Ανοικτής Εκδήλωσης
"""

import sys
sys.path.append("..")

import io
import json
import os
import re
import zipfile
from copy import deepcopy
from datetime import date
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st

from modules.database import (
    init_db, get_all_members, get_member, get_members_dropdown,
)

# ══════════════════════════════════════════════════════════════
# APP SETUP
# ══════════════════════════════════════════════════════════════
init_db()

st.set_page_config(
    page_title="Διαδικασίες & Έντυπα",
    page_icon="📋",
    layout="wide",
)

st.markdown("# 📋 Διαδικασίες & Έντυπα ΜΣΤΕ")
st.caption(
    "Επιλέξτε διαδικασία → δείτε τα βήματα → συμπληρώστε στοιχεία → "
    "λάβετε έτοιμα έντυπα για αποστολή στη Μεγ. Γεν. Γραμματεία"
)

# ══════════════════════════════════════════════════════════════
# ΣΤΑΘΕΡΕΣ — ΔΙΑΔΙΚΑΣΙΕΣ
# ══════════════════════════════════════════════════════════════
FORMS_ROOT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "forms")

DIKADIKASIEES: Dict[str, Dict] = {

    "📥 Εισδοχή Υποψηφίου": {
        "code": "eisodoxh",
        "subtitle": "Άρθρα 95–105 Γ.Κ.",
        "description": (
            "Πλήρης διαδικασία αποδοχής νέου μέλους: από την υποβολή "
            "της Συστατικής Δήλωσης ως την έγκριση της Μεγ. Στοάς."
        ),
        "deadline_note": "⏱️ Εντός 8 μηνών από έγκριση ΜΣΤ πρέπει να γίνει η μύηση (άρ. 105).",
        "steps": [
            ("1", "Απαιτούμενα έντυπα", [
                "Βιογραφικό Σημείωμα υποψηφίου",
                "Συστατική Δήλωση προτείνοντος Διδασκάλου (Αναδόχου)",
                "Αίτηση Εισδοχής υποψηφίου",
                "Υπεύθυνη Δήλωση υποψηφίου",
                "Εμπιστευτική Έκθεση εντεταλμένου Διδ (×3)",
                "2+2 φωτογραφίες (έγχρωμες, σακάκι & γραβάτα)",
            ]),
            ("2", "Υποβολή Συστατικής Δήλωσης & Βιογραφικού στη Στοά", ["Άρ. 95§1"]),
            ("3", "Ερώτηση στη Μεγ. Γεν. Γραμματεία για ιδιαίτερη κατάσταση", ["Άρ. 95§3"]),
            ("4", "Απόφαση Συμβουλίου Αξιωμ. εντός 1 μηνός", ["Άρ. 95§2 · Κ.Χ. 13"]),
            ("5", "Υποβολή Αίτησης Εισδοχής + Υπεύθυνης Δήλωσης (Σάκος Προτάσεων)", ["Άρ. 95§4 & 96"]),
            ("6", "Ανάγνωση Αίτησης & κατ' αρχήν φανερή ψηφοφορία", ["Άρ. 97§1 & 2"]),
            ("7α", "Ανάθεση σε 3 Διδ. για εμπιστ. πληρ. (εντός 7 ημ.) — έκθεση εντός 20 ημ.", ["Άρ. 97§2-4"]),
            ("7β", "Αίτηση προς λοιπές Στοές για πληροφορίες (εντός 1 μηνός)", ["Άρ. 97§4"]),
            ("7γ", "Ανάρτηση πίνακα με φωτογραφία στον Πρόναο (επί 1 μήνα)", ["Άρ. 98§1 & 2"]),
            ("8", "1 μήνα μετά: ανάγνωση 3 εκθέσεων & μυστική ψηφοφορία (σφαιρίδια)", ["Άρ. 100§1-4"]),
            ("9", "Αν αρνητικό: αποστολή ανακοίνωσης στη Μεγ. Γεν. Γραμματεία", ["Άρ. 100§3-4"]),
            ("10", "Αν θετικό: υποβολή πλήρους φακέλου στη Μεγ. Γεν. Γραμματεία", ["Άρ. 104§1"]),
            ("11", "Μεγ. Γεν. Γραμμ. εκδίδει Εγκριτικό Πίνακα Εισδοχής", []),
            ("12", "Μετά τη Μύηση: αποστολή αποκόμματος Αναγγελίας Εισδοχής", []),
        ],
        "forms": [
            {"name": "1. Βιογραφικό Σημείωμα Υποψηφίου",
             "file": "eisodoxh/1_-_ΒΙΟΓΡΑΦΙΚΟ_ΣΗΜΕΙΩΜΑ_ΥΠΟΨΗΦΙΟΥ.docx",
             "description": "Προσωπικά στοιχεία υποψηφίου"},
            {"name": "2. Συστατική Δήλωση Προτείνοντος Διδ.",
             "file": "eisodoxh/2_-_ΣΥΣΤΑΤΙΚΗ_ΔΗΛΩΣΗ_ΠΡΟΤΕΙΝΟΝΤΟΣ_ΔΙΔΑΣΚΑΛΟΥ.docx",
             "description": "Δήλωση Αναδόχου Διδασκάλου"},
            {"name": "3. Αίτηση Εισδοχής",
             "file": "eisodoxh/3_-_ΑΙΤΗΣΗ_ΕΙΣΔΟΧΗΣ__2024_05__2_.docx",
             "description": "Επίσημη αίτηση εισδοχής"},
            {"name": "4. Υπεύθυνη Δήλωση Υποψηφίου",
             "file": "eisodoxh/4_-_ΥΠΕΥΘΥΝΗ__ΔΗΛΩΣΗ__ΥΠΟΨΗΦΙΟΥ.docx",
             "description": "GDPR & δεδομένα"},
            {"name": "5. Εμπιστευτική Έκθεση Διδασκάλου",
             "file": "eisodoxh/5_-_ΕΜΠΙΣΤΕΥΤΙΚΗ_ΕΚΘΕΣΗ_ΔΙΔΑΣΚΑΛΟΥ.docx",
             "description": "Έκθεση εντεταλμένου (×3 αντίγραφα)"},
        ],
        "input_groups": [
            {
                "title": "👤 Στοιχεία Υποψηφίου",
                "fields": [
                    ("yp_eponimo",   "Επώνυμο",             "text",   ""),
                    ("yp_onoma",     "Όνομα",               "text",   ""),
                    ("yp_patronimo", "Πατρώνυμο",           "text",   ""),
                    ("yp_mitronimo", "Μητρώνυμο",           "text",   ""),
                    ("yp_gennisi",   "Ημ/νία Γέννησης",     "date",   ""),
                    ("yp_topos",     "Τόπος Γέννησης",      "text",   ""),
                    ("yp_epaggelma", "Επάγγελμα",           "text",   ""),
                    ("yp_diefthinsi","Διεύθυνση Κατοικίας", "text",   ""),
                    ("yp_poli",      "Πόλη",                "text",   "Αθήνα"),
                    ("yp_tilefono",  "Τηλέφωνο",            "text",   ""),
                    ("yp_kinito",    "Κινητό",              "text",   ""),
                    ("yp_email",     "Email",               "text",   ""),
                    ("yp_adt",       "Α.Δ.Τ. / Διαβατήριο","text",   ""),
                    ("yp_afm",       "Α.Φ.Μ.",              "text",   ""),
                    ("yp_oikogeneiaki","Οικογενειακή Κατάσταση","text",""),
                ],
            },
            {
                "title": "🤝 Στοιχεία Αναδόχου Διδασκάλου",
                "fields": [
                    ("an_eponimo",   "Επώνυμο Αναδόχου",   "member", ""),
                    ("an_onoma",     "Όνομα",               "auto",   ""),
                    ("an_arith",     "Αρ. Μητρώου Στοάς",  "auto",   ""),
                ],
            },
            {
                "title": "📋 Στοιχεία Στοάς & Διαδικασίας",
                "fields": [
                    ("st_sev",       "Σεβάσμιος",          "member", ""),
                    ("st_gramm",     "Γραμματεύς",         "member", ""),
                    ("st_imerominia_synedrias", "Ημ/νία Συνεδρίας", "date", ""),
                    ("st_arithos_protokollou",  "Αρ. Πρωτοκόλλου", "text",  ""),
                ],
            },
        ],
        "claude_hint": (
            "Πρόκειται για εισδοχή νέου μέλους σε Τεκτονική Στοά. "
            "Συμπλήρωσε τα έντυπα με επίσημη γλώσσα. "
            "Χρησιμοποίησε τεκτονικές συντομογραφίες (Σεβ∴ Διδ∴, Αδ∴ κλπ). "
            "Βάσει άρθρων 95-105 Γεν. Κανονισμού ΜΣΤΕ."
        ),
    },

    "📈 Αύξηση Μισθοδοσίας (Εταίρος)": {
        "code": "auxisi_etairos",
        "subtitle": "Άρθρο 116 Γ.Κ.",
        "description": (
            "Διαδικασία προαγωγής Μαθητή σε Εταίρο. "
            "Απαιτείται παρέλευση 1 έτους από τη μύηση και παρουσία στα 2/3 των συνεδριών."
        ),
        "deadline_note": "⏱️ Παρέλευση 1 έτους από μύηση (άρ. 113). Παρουσία σε 2/3 συνεδριών.",
        "steps": [
            ("1", "Πάροδος 1 έτους από μύηση στον βαθμό Μαθητού", ["Άρ. 113"]),
            ("2", "Παρουσία σε 2/3 συνεδριών ως Μαθητής (Βιβλίο Παρουσίας)", ["Άρ. 113"]),
            ("3", "Υποβολή Πρότασης Αύξησης Μισθοδοσίας (Σάκος Προτάσεων) — 2 Διδ. υπογράφουν", ["Άρ. 116§1"]),
            ("4", "Ανάγνωση & κατ' αρχήν φανερή ψηφοφορία", ["Άρ. 116§2"]),
            ("5", "Επόμενη συνεδρία Εταίρου: μυστική ψηφοφορία (3/4 πλειοψηφία)", ["Άρ. 116§3"]),
            ("6", "Αν αρνητικό: ανακοίνωση στη Μεγ. Γεν. Γραμματεία", []),
            ("7", "Αν θετικό: υποβολή εντύπου στη Μεγ. Γεν. Γραμματεία (πρωτότυπο + e-mail)", ["Άρ. 116§4"]),
            ("8", "Μεγ. Γεν. Γραμμ. εκδίδει Εγκριτικό Πίνακα — καταβολή μύετρων", []),
            ("9", "Μετά τη Μύηση: αποστολή αποκόμματος Αναγγελίας Αύξησης", []),
        ],
        "forms": [
            {"name": "1. Αύξηση Μισθοδοσίας Εταίρου",
             "file": "auxisi/1_-_ΑΥΞΗΣΗ_ΜΙΣΘΟΔΟΣΙΑΣ_ΕΤΑΙΡΟΥ.docx",
             "description": "Δίφυλλο έντυπο: Πρόταση + Αίτηση Έγκρισης"},
        ],
        "input_groups": [
            {
                "title": "👤 Στοιχεία Υποψηφίου",
                "fields": [
                    ("yp_melos",     "Μέλος (Υποψήφιος)",  "member", ""),
                    ("yp_eponimo",   "Επώνυμο",             "auto",   ""),
                    ("yp_onoma",     "Όνομα",               "auto",   ""),
                    ("yp_arith",     "Αρ. Μητρώου Στοάς",  "auto",   ""),
                    ("yp_arith_ms",  "Αρ. Μητρώου ΜΣ",     "auto",   ""),
                    ("yp_im_mithsis","Ημ/νία Μύησης (Μαθ.)","auto",  ""),
                    ("yp_par_synedrion","Αριθμός παρουσιών","number", ""),
                    ("yp_synolikoi","Συν. συνεδρίες Μαθητού","number",""),
                ],
            },
            {
                "title": "✍️ Προτείνοντες Διδάσκαλοι",
                "fields": [
                    ("pr1_melos",    "1ος Προτείνων",       "member", ""),
                    ("pr2_melos",    "2ος Προτείνων",       "member", ""),
                ],
            },
            {
                "title": "📋 Στοιχεία Στοάς",
                "fields": [
                    ("st_sev",       "Σεβάσμιος",          "member", ""),
                    ("st_gramm",     "Γραμματεύς",         "member", ""),
                    ("st_imerominia","Ημ/νία Συνεδρίας",   "date",   ""),
                    ("st_psifoi_yper","Ψήφοι υπέρ",        "number", ""),
                    ("st_psifoi_kata","Ψήφοι κατά",        "number", ""),
                    ("st_paronton",  "Παρόντες (αριθμός)", "number", ""),
                ],
            },
        ],
        "claude_hint": (
            "Πρόκειται για αύξηση μισθοδοσίας Μαθητή σε Εταίρο (βαθμός Εταίρου). "
            "Βάσει άρθρου 116 Γεν. Κανονισμού. Απαιτείται 1 έτος από μύηση και 2/3 παρουσίες. "
            "Συμπλήρωσε με επίσημη τεκτονική γλώσσα."
        ),
    },

    "🗳️ Εκλογές Αξιωματικών & Εγκατάσταση": {
        "code": "ekloges",
        "subtitle": "Άρθρα 73–86 Γ.Κ.",
        "description": (
            "Εκλογή Σεβασμίου, Συμβουλίου Αξιωματικών & Αντιπροσώπων. "
            "Υποβολή εντύπων στη Μεγ. Γεν. Γραμματεία εντός 10 ημερών."
        ),
        "deadline_note": "⏱️ Εντός 10 ημερών από αρχαιρεσίες υποβολή (άρ. 85§2). Εντός Οκτωβρίου.",
        "steps": [
            ("1", "Ελέγχος εκλογέων (1/3 παρουσιών τελευταίας διετίας)", ["Άρ. 73§2 & 128"]),
            ("2", "Ελέγχος εκλεξίμων (αρ. 73, 74 — ειδικές προϋποθέσεις για Σεβάσμιο)", ["Άρ. 73, 74"]),
            ("3", "Συνεδρία Μέσου Δώματος — παρουσία ≥ 1/2+1 Διδ. με δικ. ψήφου", ["Άρ. 73§3 & 138"]),
            ("4", "Εκλογή Σεβασμίου (1η ψηφοφορία)", ["Άρ. 77-83"]),
            ("5", "Εκλογή Α' & Β' Εποπτών", ["Άρ. 77"]),
            ("6", "Εκλογή λοιπών Αξιωματικών & Εξελεγκτικής Επιτροπής", ["Άρ. 77"]),
            ("7", "Τυχόν εκλογή Αντιπροσώπων (συμπίπτει ή χωριστά)", ["Άρ. 80 & 81§5"]),
            ("8", "Ανακοίνωση αποτελεσμάτων — επευφημία Στοάς", ["Άρ. 82"]),
            ("9", "Εντός 10 ημ.: υποβολή εντύπων α–ε στη Μεγ. Γεν. Γραμμ.", ["Άρ. 85§2"]),
            ("10","Εντός 15 Οκτωβρίου: Διοικητικός & Οικον. Απολογισμός", []),
            ("11","Μεγ. Στοά εγκρίνει — έκδοση Εγκριτικών Πινάκων", ["Άρ. 86"]),
            ("12","Αίτηση ημερομηνίας εγκατάστασης (εντός 30 ημ. από έγκριση)", ["Άρ. 88"]),
            ("13","Εγκατάσταση νέων Αρχών — αναγγελία στη Μεγ. Γεν. Γραμμ.", []),
        ],
        "forms": [
            {"name": "1. Αναγγελία Εκλογής Σεβασμίου & Συμβουλίου",
             "file": "ekloges/1_-_ΑΝΑΓΓΕΛΙΑ_ΕΚΛΟΓΗΣ_ΣΕΒΑΣΜΙΟΥ___ΣΥΜΒΟΥΛΙΟΥ.docx",
             "description": "Κύρια αναγγελία αποτελεσμάτων"},
            {"name": "2Α. Απόσπασμα Εκλογής Αξιωματικών",
             "file": "ekloges/2Α_-_ΑΠΟΣΠΑΣΜΑ_ΕΚΛΟΓΗΣ_ΑΞΙΩΜΑΤΙΚΩΝ.docx",
             "description": "Απόσπασμα πρακτικών συνεδρίας εκλογής"},
            {"name": "3. Αναλυτικός Πίνακας Εκλεγέντων",
             "file": "ekloges/3_-_ΑΝΑΛΥΤΙΚΟΣ_ΠΙΝΑΚΑΣ__ΕΚΛΕΓΕΝΤΩΝ_ΑΞΙΩΜΑΤΙΚΩΝ.docx",
             "description": "Πίνακας με ψήφους ανά αξίωμα"},
            {"name": "4. Κατάσταση Εκλεγέντων Αξιωματικών",
             "file": "ekloges/4_-_ΚΑΤΑΣΤΑΣΗ_ΕΚΛΕΓΕΝΤΩΝ_ΑΞΙΩΜΑΤΙΚΩΝ.docx",
             "description": "Σύνοψη εκλεγέντων"},
            {"name": "5. Ονομαστική Κατάσταση Ψηφισάντων",
             "file": "ekloges/5_-_ΟΝΟΜΑΣΤΙΚΗ_ΚΑΤΑΣΤΑΣΗ_ΨΗΦΙΣΑΝΤΩΝ.docx",
             "description": "Κατάλογος Διδ. που ψήφισαν"},
        ],
        "input_groups": [
            {
                "title": "🏛️ Νέο Συμβούλιο Αξιωματικών",
                "fields": [
                    ("ax_sev",    "Σεβάσμιος",             "member", ""),
                    ("ax_a_ep",   "Α' Επόπτης",            "member", ""),
                    ("ax_b_ep",   "Β' Επόπτης",            "member", ""),
                    ("ax_rhtor",  "Ρήτωρ",                 "member", ""),
                    ("ax_gramm",  "Γραμματεύς-Σφραγιδοφύλαξ","member",""),
                    ("ax_tamias", "Ταμίας",                "member", ""),
                    ("ax_eleon",  "Ελεονόμος",             "member", ""),
                    ("ax_tel",    "Τελετάρχης",            "member", ""),
                    ("ax_steg",   "Στεγαστής",             "member", ""),
                    ("ax_a_dok",  "Α' Δοκιμαστής",         "member", ""),
                    ("ax_b_dok",  "Β' Δοκιμαστής",         "member", ""),
                    ("ax_arxitekt","Αρχιτέκτων-Αρχιτρίκλινος","member",""),
                    ("ax_arxif",  "Αρχειοφύλαξ-Βιβλιοφύλαξ","member",""),
                    ("ax_xifok",  "Ξιφοφόρος-Σηματοφόρος",  "member", ""),
                ],
            },
            {
                "title": "➕ Πρόσθετοι Αξιωματικοί (αν υπάρχουν)",
                "fields": [
                    ("pr_rhtor",  "Πρόσθ. Ρήτωρ",          "member", ""),
                    ("pr_gramm",  "Πρόσθ. Γραμματεύς",     "member", ""),
                    ("pr_tamias", "Πρόσθ. Ταμίας",          "member", ""),
                    ("pr_eleon",  "Πρόσθ. Ελεονόμος",       "member", ""),
                    ("pr_tel",    "Πρόσθ. Τελετάρχης",      "member", ""),
                ],
            },
            {
                "title": "🗓️ Στοιχεία Αρχαιρεσιών",
                "fields": [
                    ("ekl_imerominia", "Ημ/νία Αρχαιρεσιών",   "date",   ""),
                    ("ekl_dikaiom",    "Εκλογείς (αριθμός)",    "number", ""),
                    ("ekl_paronton",   "Παρόντες κατά εκλογή",  "number", ""),
                    ("ekl_psifisantes","Ψηφίσαντες",            "number", ""),
                    ("ekl_diettia",    "Τεκτονική Διετία",      "text",   "2024-2026"),
                ],
            },
            {
                "title": "📊 Εξελεγκτική Επιτροπή",
                "fields": [
                    ("ex_1", "1ο Μέλος",  "member", ""),
                    ("ex_2", "2ο Μέλος",  "member", ""),
                    ("ex_3", "3ο Μέλος",  "member", ""),
                ],
            },
        ],
        "claude_hint": (
            "Εκλογές Αξιωματικών Τεκτονικής Στοάς. Βάσει άρθρων 73-86 Γεν. Κανονισμού ΜΣΤΕ. "
            "Η εκλογή γίνεται ανά διετία με μυστική δια ψηφοδελτίων ψηφοφορία. "
            "Χρησιμοποίησε επίσημη τεκτονική γλώσσα και συντομογραφίες."
        ),
    },

    "📣 Αναγγελία Μεταβολών": {
        "code": "metavoles",
        "subtitle": "Άρθρα 159–175 Γ.Κ.",
        "description": (
            "Αναγγελία μεταβολών μελών (θάνατος, αποχώρηση, διαγραφή, "
            "οικονομική ταξινόμηση, αναστολή κλπ.)."
        ),
        "deadline_note": "⏱️ Άμεση αναγγελία στη Μεγ. Γεν. Γραμματεία.",
        "steps": [
            ("1", "Καταγραφή μεταβολής στα πρακτικά Στοάς", []),
            ("2", "Συμπλήρωση εντύπου Αναγγελίας Μεταβολών", []),
            ("3", "Υπογραφή από Σεβ. & Γραμματεύς + Σφραγίδα Στοάς", []),
            ("4", "Αποστολή στη Μεγ. Γεν. Γραμματεία (αυτοπροσώπως ή e-mail)", []),
        ],
        "forms": [
            {"name": "1. Αναγγελία Μεταβολών",
             "file": "metavoles/1_-_ΑΝΑΓΓΕΛΙΑ_ΜΕΤΑΒΟΛΩΝ.docx",
             "description": "Έντυπο αναγγελίας μεταβολής μέλους"},
        ],
        "input_groups": [
            {
                "title": "👤 Στοιχεία Μέλους",
                "fields": [
                    ("melos",        "Μέλος",               "member", ""),
                    ("eponimo",      "Επώνυμο",             "auto",   ""),
                    ("onoma",        "Όνομα",               "auto",   ""),
                    ("arith_stoias", "Αρ. Μητρώου Στοάς",  "auto",   ""),
                    ("arith_ms",     "Αρ. Μητρώου ΜΣ",     "auto",   ""),
                    ("vathmος",      "Τεκτονικός Βαθμός",  "auto",   ""),
                ],
            },
            {
                "title": "📝 Τύπος Μεταβολής",
                "fields": [
                    ("typos_metavolis", "Τύπος Μεταβολής", "select",
                     ["Θάνατος", "Εκούσια Αποχώρηση", "Διαγραφή λόγω οφειλών",
                      "Αναστολή Δικαιωμάτων", "Επαναφορά σε Ενεργό",
                      "Αλλαγή Στοιχείων", "Άλλη μεταβολή"]),
                    ("imerominia_metavolis", "Ημ/νία Μεταβολής", "date", ""),
                    ("perigrafh",    "Περιγραφή / Παρατηρήσεις", "textarea", ""),
                ],
            },
            {
                "title": "📋 Στοιχεία Στοάς",
                "fields": [
                    ("st_sev",   "Σεβάσμιος",  "member", ""),
                    ("st_gramm", "Γραμματεύς", "member", ""),
                    ("st_date",  "Ημ/νία",     "date",   ""),
                ],
            },
        ],
        "claude_hint": (
            "Αναγγελία μεταβολής μέλους Τεκτονικής Στοάς. "
            "Συμπλήρωσε με επίσημη γλώσσα. Τεκτονικές συντομογραφίες όπου αρμόζει."
        ),
    },

    "🤝 Κοινή Συνεδρία Στοών": {
        "code": "koinh_synedria",
        "subtitle": "Εγκύκλιος ΜΣΤΕ",
        "description": (
            "Αίτηση έγκρισης κοινής συνεδρίας μεταξύ 2 ή περισσότερων Στοών "
            "υπό την Αιγίδα της ΜΣΤΕ. Υποβολή τουλάχιστον 15 ημέρες πριν."
        ),
        "deadline_note": "⏱️ Υποβολή ≥ 15 ημέρες πριν τη συνεδρία.",
        "steps": [
            ("1", "Επικοινωνία & συμφωνία με τη συνεργαζόμενη Στοά", []),
            ("2", "Συμπλήρωση Αίτησης Έγκρισης (τουλ. 15 ημ. πριν)", []),
            ("3", "Κοινοποίηση στις συνεργαζόμενες Στοές", []),
            ("4", "Αποστολή στη Μεγ. Γεν. Γραμματεία", []),
            ("5", "Μεγ. Γεν. Γραμμ. εκδίδει Εγκριτικό Πίνακα", []),
        ],
        "forms": [
            {"name": "2. Έγκριση Κοινής Συνεδρίας",
             "file": "loipa/2_-_ΕΓΚΡΙΣΗ_ΚΟΙΝΗΣ_ΣΥΝΕΔΡΙΑΣ.docx",
             "description": "Αίτηση έγκρισης κοινής συνεδρίας"},
        ],
        "input_groups": [
            {
                "title": "🏛️ Συνεργαζόμενες Στοές",
                "fields": [
                    ("stoaa_name",  "Στοά Α (η αιτούσα)",    "text", "ΑΚΡΟΠΟΛΙΣ 84"),
                    ("stoab_name",  "Στοά Β",                "text", ""),
                    ("stoac_name",  "Στοά Γ (αν υπάρχει)",   "text", ""),
                ],
            },
            {
                "title": "📅 Στοιχεία Συνεδρίας",
                "fields": [
                    ("syn_imerominia",  "Ημ/νία Κοινής Συνεδρίας", "date",   ""),
                    ("syn_ora",         "Ώρα",                     "text",   "20:00"),
                    ("syn_topos",       "Τόπος",                   "text",   "Τεκτ. Ναός"),
                    ("syn_vathmος",     "Βαθμός Εργασιών",         "select", ["Μαθητού", "Εταίρου", "Διδασκάλου"]),
                    ("syn_omilitis",    "Ομιλητής",                "member", ""),
                    ("syn_thema",       "Θέμα Ομιλίας",            "textarea",""),
                ],
            },
            {
                "title": "📋 Στοιχεία Αιτούσας Στοάς",
                "fields": [
                    ("st_sev",   "Σεβάσμιος",  "member", ""),
                    ("st_gramm", "Γραμματεύς", "member", ""),
                    ("st_date",  "Ημ/νία",     "date",   ""),
                ],
            },
        ],
        "claude_hint": (
            "Αίτηση έγκρισης κοινής συνεδρίας Τεκτονικών Στοών. "
            "Επίσημη γλώσσα, τεκτονικές συντομογραφίες. "
            "Ο ομιλητής πρέπει να αναφέρεται με αξίωμα."
        ),
    },

    "🌟 Λευκή Εορτή / Ανοικτή Εκδήλωση": {
        "code": "lefkh_eorth",
        "subtitle": "Εγκύκλιοι 89(Ο)/1999 & 24(Κ)/2002",
        "description": (
            "Αίτηση έγκρισης Λευκής Εορτής (Υιοθεσία Λυκιδέων, Μνημόσυνο, "
            "Αναγνώριση Συζυγικού Δεσμού) ή Ανοικτής Εκδήλωσης. "
            "Υποβολή ≥ 2 μήνες πριν."
        ),
        "deadline_note": "⏱️ Υποβολή ≥ 2 εργάσιμοι μήνες πριν την εκδήλωση.",
        "steps": [
            ("1", "Επιλογή τύπου εκδήλωσης (Λευκή Εορτή ή Ανοικτή Εκδήλωση)", []),
            ("2", "Τακτοποίηση οικονομικών υποχρεώσεων προς Μεγ. Θησαυροφυλάκιο", ["Εγκ. 11/1995"]),
            ("3", "Υποβολή Αίτησης Έγκρισης ≥ 2 μήνες πριν", ["Εγκ. 89(Ο)/1999"]),
            ("4", "Αν ομιλητής δεν είναι Ένδ. Αδ: υποβολή κειμένου ομιλίας για έγκριση", []),
            ("5", "Μεγ. Γεν. Γραμμ. εγκρίνει — αποστολή προσκλήσεων μόνο σε Στοές ΜΣΤΕ", ["Άρ. 189"]),
        ],
        "forms": [
            {"name": "4. Έγκριση Λευκής Εορτής / Ανοικτής Εκδήλωσης",
             "file": "loipa/4_-_ΕΓΚΡΙΣΗ_ΛΕΥΚΗΣ_ΕΟΡΤΗΣ_-_ΑΝΟΙΚ__ΕΚΔΗΛΩΣΕΩΣ.docx",
             "description": "Αίτηση έγκρισης"},
        ],
        "input_groups": [
            {
                "title": "🎉 Στοιχεία Εκδήλωσης",
                "fields": [
                    ("typos",         "Τύπος",              "select",
                     ["Υιοθεσία Λυκιδέων", "Τεκτονικό Μνημόσυνο",
                      "Τεκτ. Αναγνώριση Συζυγικού Δεσμού", "Ανοικτή Εκδήλωση"]),
                    ("imerominia",    "Ημ/νία Εκδήλωσης",  "date",   ""),
                    ("ora",           "Ώρα",                "text",   "19:30"),
                    ("topos",         "Τόπος",              "text",   ""),
                    ("omilitis",      "Ομιλητής",           "member", ""),
                    ("thema_omilias", "Θέμα Ομιλίας",      "textarea",""),
                    ("ypeuthinos",    "Υπεύθυνος Οργ.",     "member", ""),
                ],
            },
            {
                "title": "📋 Στοιχεία Στοάς",
                "fields": [
                    ("st_sev",   "Σεβάσμιος",  "member", ""),
                    ("st_gramm", "Γραμματεύς", "member", ""),
                    ("st_date",  "Ημ/νία",     "date",   ""),
                ],
            },
        ],
        "claude_hint": (
            "Αίτηση έγκρισης Λευκής Εορτής ή Ανοικτής Εκδήλωσης Τεκτονικής Στοάς. "
            "Βάσει εγκυκλίων 89(Ο)/1999 και 24(Κ)/2002 ΜΣΤΕ. "
            "Επίσημη γλώσσα και τεκτονικές συντομογραφίες."
        ),
    },
}

STOAA_NAME = "ΑΚΡΟΠΟΛΙΣ"
STOAA_NUMBER = "84"
STOAA_ANATOLI = "Αθηνών"

# ══════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════
def get_anthropic_key() -> str:
    try:
        return (
            st.secrets.get("AI", {}).get("ANTHROPIC_API_KEY")
            or st.secrets.get("ANTHROPIC_API_KEY", "")
        )
    except Exception:
        return ""


def load_members_dict() -> Dict[int, Dict]:
    """Returns {id: member_dict}"""
    try:
        df = get_all_members()
        result = {}
        for _, row in df.iterrows():
            result[int(row["id"])] = row.to_dict()
        return result
    except Exception:
        return {}


def member_display_name(m: Dict) -> str:
    vatm = m.get("τεκτονικός_βαθμός", "")
    prefix = {"Μαθητής": "Αδ∴", "Εταίρος": "Αδ∴", "Διδάσκαλος": "Αδ∴"}.get(vatm, "Αδ∴")
    return f"{prefix} {m.get('επώνυμο','')} {m.get('όνομα','')}".strip()


def read_docx_text(path: str) -> str:
    """Extract raw text from docx for passing to Claude."""
    try:
        from docx import Document
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                seen = set()
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid not in seen and cell.text.strip():
                        lines.append(cell.text.strip())
                        seen.add(cid)
        return "\n".join(lines)
    except Exception as e:
        return f"[Σφάλμα ανάγνωσης: {e}]"


# ══════════════════════════════════════════════════════════════
# ΑΝΤΙΣΤΟΙΧΙΑ ΑΞΙΩΜΑΤΩΝ → ΠΕΔΙΑ ΧΡΗΣΤΗ
# ══════════════════════════════════════════════════════════════

# Αντιστοιχία κειμένου στήλης ΑΞΙΩΜΑ → key στα collected data
AXIOMA_TO_KEY: Dict[str, str] = {
    # Τακτικοί
    "ΣΕΒΑΣΜΙΟΣ":              "ax_sev",
    "Α΄ ΕΠΟΠΤΗΣ":             "ax_a_ep",
    "Β΄ ΕΠΟΠΤΗΣ":             "ax_b_ep",
    "ΡΗΤΩΡ":                  "ax_rhtor",
    "ΓΡΑΜΜ. - ΣΦΡΑΓΙΔ.":     "ax_gramm",
    "Α΄ ΔΟΚΙΜΑΣΤΗΣ":          "ax_a_dok",
    "ΤΑΜΙΑΣ":                 "ax_tamias",
    "ΕΛΕΟΝΟΜΟΣ":              "ax_eleon",
    "ΤΕΛΕΤΑΡΧΗΣ":             "ax_tel",
    "ΣΤΕΓΑΣΤΗΣ":              "ax_steg",
    "Β΄ ΔΟΚΙΜΑΣΤΗΣ":          "ax_b_dok",
    "ΑΡΧΙΤ. - ΑΡΧΙΤΡ.":      "ax_arxitekt",
    "ΑΡΧΙΤ.  ΑΡΧΙΤΡ.":       "ax_arxitekt",
    "ΑΡΧΕΙΟΦ. - ΒΙΒΛΙΟΦ.":   "ax_arxif",
    "ΞΙΦΟΦ. - ΣΗΜΑΙΟΦ.":     "ax_xifok",
    # Πρόσθετοι — με prefix "ΙΙ" section
    "ΡΗΤΩΡ_Π":                "pr_rhtor",
    "ΓΡΑΜΜ. - ΣΦΡΑΓΙΔ._Π":   "pr_gramm",
    "ΤΑΜΙΑΣ_Π":               "pr_tamias",
    "ΕΛΕΟΝΟΜΟΣ_Π":            "pr_eleon",
    "ΤΕΛΕΤΑΡΧΗΣ_Π":           "pr_tel",
    # Εξελεγκτική
    "ΜΕΛΟΣ 1":                "ex_1",
    "ΜΕΛΟΣ 2":                "ex_2",
    "ΜΕΛΟΣ 3":                "ex_3",
    # Απόσπασμα — για το αξίωμα του:
    "Ρήτορα":                 "ax_rhtor",
    "Γραμματέα - Σφραγιδοφύλακα": "ax_gramm",
    "Α΄ Δοκιμαστή":           "ax_a_dok",
    "Ταμία":                  "ax_tamias",
    "Ελεονόμου":              "ax_eleon",
    "Τελετάρχη":              "ax_tel",
    "Στεγαστή":               "ax_steg",
    "Β΄ Δοκιμαστή":           "ax_b_dok",
    "Αρχιτέκτονα - Αρχιτρίκλινου": "ax_arxitekt",
    "Αρχειοφύλακα - Βιβλιοφύλακα": "ax_arxif",
    "Ξιφοφόρου - Σημαιοφόρου":     "ax_xifok",
    "Πρόσθετου Ρήτορα":       "pr_rhtor",
    "Πρόσθετου Γραμματέα - Σφραγιδοφύλακα": "pr_gramm",
    "Πρόσθ. Γραμμ - Σφραγιδοφ":    "pr_gramm",
    "Πρόσθετου Ταμία":        "pr_tamias",
    "Πρόσθετου Ελεονόμου":    "pr_eleon",
    "Πρόσθετου Τελετάρχη":    "pr_tel",
    "Μέλους της Εξελεγκτικής Επιτροπής": "ex_1",
    "Μέλους Εξελεγκτικής Επιτροπής":     "ex_1",
    # Α'/Β' Επόπτης
    "Α΄ Επόπτου":             "ax_a_ep",
    "Β΄ Επόπτου":             "ax_b_ep",
}

# ══════════════════════════════════════════════════════════════
# ΒΟΗΘΗΤΙΚΕΣ ΣΥΝΑΡΤΗΣΕΙΣ ΕΓΓΡΑΦΗΣ ΣΤΟΥΣ ΠΙΝΑΚΕΣ
# ══════════════════════════════════════════════════════════════

def _unique_cells(row) -> List:
    """Επιστρέφει μόνο τα μοναδικά cells μιας γραμμής (χωρίς duplicates από merged cells)."""
    seen = set()
    cells = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            cells.append(cell)
            seen.add(cid)
    return cells


def _write_cell(cell, text: str) -> None:
    """Γράφει κείμενο σε ένα cell διατηρώντας τη μορφοποίηση (font, size κλπ)."""
    text = str(text or "").strip()
    if not text:
        return
    for para in cell.paragraphs:
        if para.runs:
            # Κρατάει τη μορφοποίηση του πρώτου run
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""
            return
    # Αν δεν υπάρχουν runs, προσθέτουμε νέο
    if cell.paragraphs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)


def _replace_in_para(para, old: str, new: str) -> bool:
    """Αντικαθιστά κείμενο σε paragraph που μπορεί να είναι split σε πολλά runs."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_text = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    return True


def _member_full_name(m: Optional[Dict]) -> str:
    if not m:
        return ""
    return f"{m.get('επώνυμο', '')} {m.get('όνομα', '')}".strip()


def _member_full_with_patronymic(m: Optional[Dict]) -> str:
    if not m:
        return ""
    parts = [m.get('επώνυμο', ''), m.get('όνομα', '')]
    pat = m.get('πατρώνυμο', '')
    if pat:
        parts.append(f"του {pat}")
    return " ".join(p for p in parts if p).strip()


# ══════════════════════════════════════════════════════════════
# ΚΥΡΙΑ ΣΥΝΑΡΤΗΣΗ ΕΓΓΡΑΦΗΣ — JSON ΔΟΜΗ → ΑΜΕΣΗ ΕΓΓΡΑΦΗ
# ══════════════════════════════════════════════════════════════

def fill_docx_smart(
    template_path: str,
    form_file_key: str,
    collected: Dict[str, Any],
    members_by_field: Dict[str, Optional[Dict]],
    stoaa_data: Dict[str, str],
) -> bytes:
    """
    Κύρια συνάρτηση συμπλήρωσης εντύπων.
    Χρησιμοποιεί άμεση εγγραφή στα σωστά cells βάσει δομής κάθε εντύπου.
    ΔΕΝ κάνει text replacement — γράφει απευθείας στα κενά cells.
    """
    from docx import Document

    doc = Document(template_path)
    fname = os.path.basename(form_file_key).upper()

    # ════════════════════════════════════════════════════════
    # DEBUG PANEL
    # ════════════════════════════════════════════════════════
    with st.expander(f"🔍 DEBUG: {os.path.basename(form_file_key)}", expanded=True):
        st.markdown(f"**form_file_key:** `{form_file_key}`")
        st.markdown(f"**fname (uppercase):** `{fname}`")

        st.markdown("**Έλεγχος συνθήκης εντύπου:**")
        st.write({
            "4_ΚΑΤΑΣΤΑΣΗ match": "4_-_ΚΑΤΑΣΤΑΣΗ" in fname or "ΚΑΤΑΣΤΑΣΗ_ΕΚΛ" in fname,
            "3_ΑΝΑΛΥΤΙΚΟΣ match": "3_-_ΑΝΑΛΥΤΙΚ" in fname or "ΑΝΑΛΥΤΙΚΟΣ_ΠΙΝ" in fname,
            "2Α_ΑΠΟΣΠΑΣΜΑ match": "2" in fname and "ΑΠΟΣΠ" in fname,
            "1_ΑΝΑΓΓΕΛΙΑ match": "1_-_ΑΝΑΓΓΕΛΙΑ_ΕΚΛΟΓ" in fname or "ΑΝΑΓΓΕΛΙΑ_ΕΚΛΟΓ" in fname,
            "5_ΟΝΟΜΑΣΤΙΚΗ match": "5_-_ΟΝΟΜΑΣΤ" in fname or "ΟΝΟΜΑΣΤΙΚΗ_ΚΑΤΑΣ" in fname,
        })

        st.markdown(f"**members_by_field keys:** `{list(members_by_field.keys())}`")
        st.markdown(f"**Συνολικά members:** {sum(1 for v in members_by_field.values() if v is not None)}")

        filled_members = {k: f"{v.get('επώνυμο','')} {v.get('όνομα','')}" 
                         for k, v in members_by_field.items() if v}
        if filled_members:
            st.markdown("**Μέλη που επιλέχθηκαν:**")
            for k, v in filled_members.items():
                st.write(f"  - `{k}` → {v}")
        else:
            st.error("⚠️ ΔΕΝ ΕΠΙΛΕΧΘΗΚΑΝ ΜΕΛΗ! Πήγαινε στη καρτέλα 'Συμπλήρωση' και επίλεξε αξιωματικούς.")

        st.markdown(f"**collected keys:** `{list(collected.keys())}`")
        st.markdown(f"**ekl_diettia:** `{collected.get('ekl_diettia', 'ΔΕΝ ΒΡΕΘΗΚΕ')}`")
        st.markdown(f"**ekl_imerominia:** `{collected.get('ekl_imerominia', 'ΔΕΝ ΒΡΕΘΗΚΕ')}`")
    # ════════════════════════════════════════════════════════

    # ── Βοηθητική: παίρνει μέλος από πεδίο ──────────────────
    def get_member(field_key: str) -> Optional[Dict]:
        return members_by_field.get(field_key)

    def get_name(field_key: str) -> str:
        m = get_member(field_key)
        return _member_full_name(m)

    def get_name_pat(field_key: str) -> str:
        m = get_member(field_key)
        return _member_full_with_patronymic(m)

    def get_surname(field_key: str) -> str:
        m = get_member(field_key)
        return m.get('επώνυμο', '') if m else ''

    def get_firstname(field_key: str) -> str:
        m = get_member(field_key)
        return m.get('όνομα', '') if m else ''

    def get_field(fid: str) -> str:
        return str(collected.get(fid, '') or '')

    diettia = get_field("ekl_diettia") or "2024-2026"
    diettia_parts = diettia.replace("-", " / ").replace("–", " / ")
    imerominia = get_field("ekl_imerominia")
    psifisantes = get_field("ekl_psifisantes")

    # ══════════════════════════════════════════════════════════
    # FORM 4: ΚΑΤΑΣΤΑΣΗ ΕΚΛΕΓΕΝΤΩΝ ΑΞΙΩΜΑΤΙΚΩΝ
    # Δομή: 1 table, 32 rows, 12 unique cols per data row
    # Col 3: ΑΞΙΩΜΑ, Col 4: ΕΠΩΝΥΜΟ, Col 5: ΟΝΟΜΑ,
    # Col 6: ΔΙΕΥΘΥΝΣΗ, Col 7: ΠΟΛΗ-ΤΚ, Col 8: ΚΙΝΗΤΟ,
    # Col 9: ΣΤΑΘΕΡΟ, Col 10: EMAIL
    # ══════════════════════════════════════════════════════════
    if "4_-_ΚΑΤΑΣΤΑΣΗ" in fname or "ΚΑΤΑΣΤΑΣΗ_ΕΚΛ" in fname:
        table = doc.tables[0]
        in_prostheti = False
        filled_count = 0

        # Header cells
        for r_i, row in enumerate(table.rows):
            cells = _unique_cells(row)
            if len(cells) < 2:
                continue
            for c_i, cell in enumerate(cells):
                ct = cell.text.strip()
                if "Τεκτ" in ct and "Διετία" in ct and c_i + 1 < len(cells):
                    _write_cell(cells[c_i + 1], diettia_parts)
                elif ct in ("Σ Στοά:", "Σ∴ Στ∴:", "Σ. Στ.:") and c_i + 1 < len(cells):
                    _write_cell(cells[c_i + 1], f"{stoaa_data['name']}")
                elif ct in ("Εν Αν:", "Εν Αν∴:") and c_i + 1 < len(cells):
                    _write_cell(cells[c_i + 1], stoaa_data['anatoli'])

        # Official rows
        debug_rows = []
        for r_i, row in enumerate(table.rows):
            cells = _unique_cells(row)
            if len(cells) < 5:
                continue
            axioma = cells[3].text.strip() if len(cells) > 3 else ""

            if "ΠΡΟΣΘΕΤΟΙ" in axioma:
                in_prostheti = True
                continue
            if axioma in ("Ι", "ΙΙ", "ΙΙΙ"):
                in_prostheti = False
                continue

            lookup = axioma + ("_Π" if in_prostheti else "")
            field_key = AXIOMA_TO_KEY.get(lookup) or AXIOMA_TO_KEY.get(axioma)
            if not field_key or not axioma or axioma.startswith("Α/Α"):
                continue

            m = get_member(field_key)
            debug_rows.append({
                "row": r_i,
                "axioma": axioma,
                "field_key": field_key,
                "member_found": m is not None,
                "name": f"{m.get('επώνυμο','')} {m.get('όνομα','')}" if m else "—"
            })

            if not m:
                continue

            if len(cells) > 4:
                _write_cell(cells[4], m.get('επώνυμο', ''))
            if len(cells) > 5:
                _write_cell(cells[5], m.get('όνομα', ''))
            if len(cells) > 6:
                _write_cell(cells[6], m.get('διεύθυνση', ''))
            if len(cells) > 7:
                poli = m.get('πόλη', '')
                tk = m.get('τ_κ', '') or m.get('τκ', '')
                _write_cell(cells[7], f"{poli} {tk}".strip())
            if len(cells) > 8:
                _write_cell(cells[8], m.get('κινητό', '') or m.get('κινητο', ''))
            if len(cells) > 9:
                _write_cell(cells[9], m.get('τηλέφωνο', '') or m.get('τηλεφωνο', ''))
            if len(cells) > 10:
                _write_cell(cells[10], m.get('email', ''))
            filled_count += 1

        with st.expander("🔍 DEBUG Form 4 — Γραμμές πίνακα", expanded=True):
            st.markdown(f"**Συνολικά γραμμές αξιωματικών που βρέθηκαν:** {len(debug_rows)}")
            st.markdown(f"**Γραμμές που συμπληρώθηκαν:** {filled_count}")
            for dr in debug_rows:
                icon = "✅" if dr["member_found"] else "❌"
                st.write(f"{icon} R{dr['row']} `{dr['axioma']}` → key:`{dr['field_key']}` → {dr['name']}")

    # ══════════════════════════════════════════════════════════
    # FORM 3: ΑΝΑΛΥΤΙΚΟΣ ΠΙΝΑΚΑΣ ΕΚΛΕΓΕΝΤΩΝ
    # Δομή: 1 table, 37 rows
    # Col 4: ΟΝΟΜΑΤΕΠΩΝΥΜΟ & ΠΑΤΡΩΝΥΜΟ, Col 5: Α.Μ.Μ.Σ.
    # Col 6: ΜΑΘ., Col 7: ΔΙΔ., Col 8: ΜΕΛΟΣ ΑΠΟ, Col 9: ΛΕΥΚ
    # ══════════════════════════════════════════════════════════
    elif "3_-_ΑΝΑΛΥΤΙΚ" in fname or "ΑΝΑΛΥΤΙΚΟΣ_ΠΙΝ" in fname:
        table = doc.tables[0]
        in_prostheti = False
        ex_count = 0  # counter for ΕΞΕΛΕΓΚΤΙΚΗ members

        for r_i, row in enumerate(table.rows):
            cells = _unique_cells(row)
            if len(cells) < 5:
                continue
            axioma = cells[3].text.strip() if len(cells) > 3 else ""

            # Header cells
            for c_i, cell in enumerate(cells):
                ct = cell.text.strip()
                if "Τεκτονική Διετία" in ct:
                    _replace_in_para(cell.paragraphs[0] if cell.paragraphs else None.__class__, "20", diettia_parts) if False else None
                    _write_cell(cells[c_i], f"Τεκτονική Διετία : {diettia_parts}")
                elif "Ημερομηνία αρχαιρεσιών" in ct and imerominia and c_i + 1 < len(cells):
                    _write_cell(cells[c_i + 1], imerominia)
                elif "Αριθμός ψηφισάντων" in ct and psifisantes and c_i + 1 < len(cells):
                    _write_cell(cells[c_i + 1], psifisantes)

            # Section markers
            if "ΠΡΟΣΘΕΤΟΙ" in axioma:
                in_prostheti = True
                continue
            if "ΕΞΕΛΕΓΚΤ" in axioma:
                in_prostheti = False
                ex_count = 0
                continue

            # Εξελεγκτική επιτροπή
            if axioma.startswith("ΜΕΛΟΣ"):
                ex_count += 1
                field_key = f"ex_{ex_count}"
            else:
                lookup = axioma + ("_Π" if in_prostheti else "")
                field_key = AXIOMA_TO_KEY.get(lookup) or AXIOMA_TO_KEY.get(axioma)

            if not field_key or not axioma or axioma.startswith("Α/Α"):
                continue

            m = get_member(field_key)
            if not m:
                continue

            full_name = _member_full_with_patronymic(m)
            amms = str(m.get('αρ_μητρώου_μσ', '') or m.get('arith_ms', '') or '')
            im_math = str(m.get('ημ_μύησης', '') or '')
            im_did = str(m.get('ημ_μύησης_διδ', '') or m.get('ημ_αύξησης_διδ', '') or '')
            melos_apo = str(m.get('μέλος_από', '') or m.get('melos_apo', '') or '')

            if len(cells) > 4:
                _write_cell(cells[4], full_name)
            if len(cells) > 5:
                _write_cell(cells[5], amms)
            if len(cells) > 6:
                _write_cell(cells[6], im_math)
            if len(cells) > 7:
                _write_cell(cells[7], im_did)
            if len(cells) > 8:
                _write_cell(cells[8], melos_apo)

    # ══════════════════════════════════════════════════════════
    # FORM 2Α: ΑΠΟΣΠΑΣΜΑ ΕΚΛΟΓΗΣ ΑΞΙΩΜΑΤΙΚΩΝ
    # 8 ξεχωριστοί πίνακες
    # ══════════════════════════════════════════════════════════
    elif "2" in fname and "ΑΠΟΣΠ" in fname:
        # Paragraphs με tabs — αντικαθιστούμε τα κενά/tabs
        for para in doc.paragraphs:
            ft = "".join(r.text for r in para.runs)
            if "Αρχαιρεσίες" in ft and "Οκτωβρίου" in ft:
                # Συμπλήρωση περιόδου αρχαιρεσιών
                parts = diettia.split("-") if "-" in diettia else diettia.split("/")
                if len(parts) == 2:
                    new = ft
                    for r in para.runs:
                        if "Οκτωβρίου 20" in r.text:
                            r.text = r.text.replace("20    ", f"20{parts[1][-2:]}")
                        if "Σεπτεμβρίου 20" in r.text:
                            r.text = r.text.replace("20    ", f"20{parts[0][-2:]}")

        # Table 0: Επιτροπή καταμέτρησης [α/α, Ονοματεπώνυμο, Αξίωμα]
        # Γράψε πρώην Σεβ, Ρήτορα, Γραμματέα, Α' Δοκιμαστή
        epitropi_map = {
            "Πρώην Σεβάσμιος": "ax_sev",  # ο απερχόμενος Σεβ
            "Ρήτορας":          "ax_rhtor",
            "Γραμματέας - Σφραγιδοφύλακας": "ax_gramm",
            "Α΄ Δοκιμαστής":   "ax_a_dok",
        }
        if len(doc.tables) > 0:
            for row in doc.tables[0].rows[1:]:
                cells = _unique_cells(row)
                if len(cells) >= 3:
                    axioma_text = cells[2].text.strip()
                    fk = epitropi_map.get(axioma_text)
                    if fk:
                        _write_cell(cells[1], get_name(fk))

        # Table 1: Ψηφολέκτες [α/α, Ονοματεπώνυμο, Αξίωμα]
        psifol_map = {"Τελετάρχης": "ax_tel"}
        if len(doc.tables) > 1:
            for row in doc.tables[1].rows[1:]:
                cells = _unique_cells(row)
                if len(cells) >= 3:
                    fk = psifol_map.get(cells[2].text.strip())
                    if fk:
                        _write_cell(cells[1], get_name(fk))

        # Table 3: Υποψήφιοι Α'/Β' Εποπτών [α/α, Αξίωμα, Ονοματεπώνυμο]
        # Βάζουμε τον εκλεγέντα στη θέση 1 (πρώτη υποψηφιότητα)
        if len(doc.tables) > 3:
            ep_rows = doc.tables[3].rows[1:]
            for row in ep_rows:
                cells = _unique_cells(row)
                if len(cells) >= 3:
                    ax = cells[1].text.strip()
                    fk = AXIOMA_TO_KEY.get(ax)
                    if fk and not cells[2].text.strip():
                        _write_cell(cells[2], get_name(fk))

        # Table 4: Εκλεγέντες Α'/Β' Εποπτών [α/α, Αξίωμα, Ονοματεπώνυμο, Λ. Ψήφοι]
        if len(doc.tables) > 4:
            for row in doc.tables[4].rows[1:]:
                cells = _unique_cells(row)
                if len(cells) >= 3:
                    ax = cells[1].text.strip()
                    fk = AXIOMA_TO_KEY.get(ax)
                    if fk:
                        _write_cell(cells[2], get_name(fk))

        # Table 5: Υποψήφιοι λοιπών [α/α, Αξίωμα, Ονοματεπώνυμο]
        if len(doc.tables) > 5:
            for row in doc.tables[5].rows[1:]:
                cells = _unique_cells(row)
                if len(cells) >= 3:
                    ax = cells[1].text.strip()
                    fk = AXIOMA_TO_KEY.get(ax)
                    if fk and not cells[2].text.strip():
                        _write_cell(cells[2], get_name(fk))

        # Table 6: Εκλεγέντες λοιπών [α/α, Αξίωμα, Ονοματεπώνυμο, Λ. Ψήφοι]
        if len(doc.tables) > 6:
            for row in doc.tables[6].rows[1:]:
                cells = _unique_cells(row)
                if len(cells) >= 3:
                    ax = cells[1].text.strip()
                    fk = AXIOMA_TO_KEY.get(ax)
                    if fk:
                        _write_cell(cells[2], get_name(fk))

    # ══════════════════════════════════════════════════════════
    # FORM 1: ΑΝΑΓΓΕΛΙΑ ΕΚΛΟΓΗΣ ΣΕΒΑΣΜΙΟΥ
    # Παραγράφους με κενά/dashes — text replacement
    # ══════════════════════════════════════════════════════════
    elif "1_-_ΑΝΑΓΓΕΛΙΑ_ΕΚΛΟΓ" in fname or "ΑΝΑΓΓΕΛΙΑ_ΕΚΛΟΓ" in fname:
        sev_name = get_name("ax_sev")
        gramm_name = get_name("ax_gramm")

        for para in doc.paragraphs:
            ft = "".join(r.text for r in para.runs)
            if "27/04/2026" in ft:
                new_ft = ft.replace("27/04/2026", imerominia or "27/04/2026")
                if para.runs:
                    para.runs[0].text = new_ft
                    for r in para.runs[1:]:
                        r.text = ""

        # Header table
        for table in doc.tables:
            for row in table.rows:
                cells = _unique_cells(row)
                for c_i, cell in enumerate(cells):
                    ct = cell.text.strip()
                    if "ΜΕΓ. ΓΕΝ. ΓΡΑΜΜΑΤΕΙΑ" in ct or "Ημερομηνία" in ct:
                        if c_i + 1 < len(cells):
                            _write_cell(cells[c_i + 1], imerominia)

    # ══════════════════════════════════════════════════════════
    # FORM 5: ΟΝΟΜΑΣΤΙΚΗ ΚΑΤΑΣΤΑΣΗ ΨΗΦΙΣΑΝΤΩΝ
    # Header + κενές γραμμές πίνακα — συμπληρώνει από μέλη
    # ══════════════════════════════════════════════════════════
    elif "5_-_ΟΝΟΜΑΣΤ" in fname or "ΟΝΟΜΑΣΤΙΚΗ_ΚΑΤΑΣ" in fname:
        if doc.tables:
            table = doc.tables[0]
            # Header info
            for r_i, row in enumerate(table.rows):
                cells = _unique_cells(row)
                for c_i, cell in enumerate(cells):
                    ct = cell.text.strip()
                    if "Αρχαιρεσίες" in ct and c_i + 1 < len(cells):
                        _write_cell(cells[c_i + 1], imerominia)
                    elif "Ανατολή" in ct and c_i + 1 < len(cells):
                        _write_cell(cells[c_i + 1], stoaa_data.get('anatoli', ''))

    # ══════════════════════════════════════════════════════════
    # GENERIC FALLBACK — Claude για τα υπόλοιπα έντυπα
    # ══════════════════════════════════════════════════════════
    else:
        # Για μη-εκλογικά έντυπα: χρησιμοποιεί Claude με JSON approach
        return _claude_fill_generic(template_path, collected, members_by_field, stoaa_data)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _claude_fill_generic(
    template_path: str,
    collected: Dict[str, Any],
    members_by_field: Dict[str, Optional[Dict]],
    stoaa_data: Dict[str, str],
) -> bytes:
    """
    Για μη-εκλογικά έντυπα: Claude επιστρέφει JSON με
    {table_idx, row_idx, cell_idx, value} και γράφουμε άμεσα.
    """
    from docx import Document

    key = get_anthropic_key()
    if not key:
        # Επιστρέφει το αρχείο αμετάβλητο
        with open(template_path, "rb") as f:
            return f.read()

    # 1. Κάνε docx → JSON δομή
    doc = Document(template_path)
    doc_structure = {"paragraphs": [], "tables": []}

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            doc_structure["paragraphs"].append({"idx": i, "text": para.text[:200]})

    for t_i, table in enumerate(doc.tables):
        table_data = {"table_idx": t_i, "rows": []}
        for r_i, row in enumerate(table.rows):
            cells_data = []
            for c_i, cell in enumerate(_unique_cells(row)):
                cells_data.append({"cell_idx": c_i, "text": cell.text.strip()[:80]})
            table_data["rows"].append({"row_idx": r_i, "cells": cells_data})
        doc_structure["tables"].append(table_data)

    # 2. Συλλογή user data
    user_data_flat = {k: str(v or "") for k, v in collected.items()}
    for field_key, m in members_by_field.items():
        if m:
            user_data_flat[f"{field_key}_full"] = _member_full_name(m)
            user_data_flat[f"{field_key}_surname"] = m.get('επώνυμο', '')
            user_data_flat[f"{field_key}_name"] = m.get('όνομα', '')

    # 3. Claude fills JSON
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=key)

        system = """
Είσαι Γραμματεύς Τεκτονικής Στοάς. Σου δίνεται η JSON δομή ενός εντύπου ΜΣΤΕ
(paragraphs και tables με ακριβείς συντεταγμένες κελιών) και τα στοιχεία του χρήστη.

Επέστρεψε JSON array με fill instructions:
[
  {"type": "table_cell", "table_idx": 0, "row_idx": 2, "cell_idx": 3, "value": "ΠΑΡΙΣΗΣ ΓΕΩΡΓΙΟΣ"},
  {"type": "paragraph", "para_idx": 5, "old_text": "........", "new_text": "27/04/2026"}
]

Κανόνες:
- Γράψε ΜΟΝΟ σε κενά cells ή cells με placeholder (......, tabs)
- Χρησιμοποίησε επίσημη τεκτονική γλώσσα
- Μόνο valid JSON array, χωρίς markdown
""".strip()

        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=3000,
            system=system,
            messages=[{"role": "user", "content":
                f"Στοιχεία Στοάς: {stoaa_data}\n\n"
                f"Δεδομένα χρήστη:\n{json.dumps(user_data_flat, ensure_ascii=False)[:2000]}\n\n"
                f"Δομή εντύπου:\n{json.dumps(doc_structure, ensure_ascii=False)[:4000]}"
            }],
        )
        raw = msg.content[0].text if msg.content else "[]"
        raw = re.sub(r"^```(?:json)?\s*", "", raw.strip(), flags=re.MULTILINE)
        raw = re.sub(r"\s*```\s*$", "", raw, flags=re.MULTILINE).strip()
        raw = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\t]", " ", raw)
        fills = json.loads(raw)

        # 4. Apply fills
        for fill in fills:
            try:
                if fill.get("type") == "table_cell":
                    table = doc.tables[fill["table_idx"]]
                    row = table.rows[fill["row_idx"]]
                    cell = _unique_cells(row)[fill["cell_idx"]]
                    _write_cell(cell, fill["value"])
                elif fill.get("type") == "paragraph":
                    para = doc.paragraphs[fill["para_idx"]]
                    _replace_in_para(para, fill.get("old_text", ""), fill["new_text"])
            except (IndexError, KeyError, Exception):
                pass

    except Exception as e:
        st.warning(f"⚠️ Claude error: {e}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def create_zip(files: List[Tuple[str, bytes]]) -> bytes:
    """Create a zip with multiple docx files."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files:
            zf.writestr(name, data)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════
# UI — ΕΠΙΛΟΓΗ ΔΙΑΔΙΚΑΣΙΑΣ
# ══════════════════════════════════════════════════════════════
col_left, col_right = st.columns([1, 2])

with col_left:
    st.markdown("### 📌 Επιλογή Διαδικασίας")
    selected = st.selectbox(
        "Διαδικασία",
        list(DIKADIKASIEES.keys()),
        label_visibility="collapsed",
    )

proc = DIKADIKASIEES[selected]

with col_right:
    st.info(f"**{selected}** — {proc['subtitle']}\n\n{proc['description']}")
    if proc.get("deadline_note"):
        st.warning(proc["deadline_note"])

st.markdown("---")

# ══════════════════════════════════════════════════════════════
# ΚΑΡΤΕΛΕΣ: Βήματα | Έντυπα | Συμπλήρωση
# ══════════════════════════════════════════════════════════════
tab_steps, tab_forms, tab_fill = st.tabs([
    "📋 Βήματα Διαδικασίας",
    "📁 Απαιτούμενα Έντυπα",
    "✍️ Συμπλήρωση & Λήψη",
])

# ──────────────────────────────────────────────────────────────
# TAB 1: ΒΗΜΑΤΑ
# ──────────────────────────────────────────────────────────────
with tab_steps:
    st.markdown(f"### Βήματα Διαδικασίας: {selected}")
    st.caption(
        "Βάσει Οδηγού Διαδικασιών & Απαιτούμενων Ενεργειών (Δεκ. 2017) "
        "και Γεν. Κανονισμού ΜΣΤΕ"
    )

    for step_num, step_title, refs in proc["steps"]:
        ref_str = " · ".join(refs) if refs else ""
        with st.expander(f"**Βήμα {step_num}** — {step_title}", expanded=False):
            if isinstance(refs, list) and refs and refs[0].startswith("•"):
                for r in refs:
                    st.markdown(r)
            elif ref_str:
                st.caption(f"📖 {ref_str}")
            else:
                st.caption("Εσωτερική διαδικασία Στοάς")

    # Απαιτούμενα έντυπα summary
    st.markdown("---")
    st.markdown("#### 📎 Απαιτούμενα Έντυπα")
    for f in proc["forms"]:
        path = os.path.join(FORMS_ROOT, f["file"])
        exists = os.path.exists(path)
        icon = "✅" if exists else "❌"
        st.markdown(f"{icon} **{f['name']}** — {f['description']}")

# ──────────────────────────────────────────────────────────────
# TAB 2: ΕΝΤΥΠΑ (preview + download blank)
# ──────────────────────────────────────────────────────────────
with tab_forms:
    st.markdown(f"### Έντυπα: {selected}")
    st.caption("Κατεβάστε τα κενά πρότυπα ή δείτε το περιεχόμενό τους.")

    all_blank = []
    for f in proc["forms"]:
        path = os.path.join(FORMS_ROOT, f["file"])
        st.markdown(f"#### {f['name']}")
        st.caption(f['description'])

        if os.path.exists(path):
            with open(path, "rb") as fh:
                data = fh.read()

            col1, col2 = st.columns([1, 3])
            with col1:
                fname = os.path.basename(f["file"])
                st.download_button(
                    f"⬇️ Κενό έντυπο",
                    data=data,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"blank_{f['file']}",
                    use_container_width=True,
                )
                all_blank.append((fname, data))
            with col2:
                with st.expander("👁️ Προβολή κειμένου εντύπου"):
                    text = read_docx_text(path)
                    st.text(text[:1500] + ("…" if len(text) > 1500 else ""))
        else:
            st.error(f"❌ Δεν βρέθηκε: {path}")

        st.markdown("---")

    if len(all_blank) > 1:
        zip_data = create_zip(all_blank)
        st.download_button(
            f"📦 Λήψη όλων των κενών εντύπων (.zip)",
            data=zip_data,
            file_name=f"εντυπα_{proc['code']}.zip",
            mime="application/zip",
            use_container_width=True,
        )

# ──────────────────────────────────────────────────────────────
# TAB 3: ΣΥΜΠΛΗΡΩΣΗ
# ──────────────────────────────────────────────────────────────
with tab_fill:
    st.markdown(f"### ✍️ Συμπλήρωση Εντύπων: {selected}")
    st.caption(
        "Εισάγετε τα στοιχεία παρακάτω. Το σύστημα θα προσυμπληρώσει από "
        "το Μητρώο όπου είναι δυνατό και θα χρησιμοποιήσει Claude για "
        "αυτόματη συμπλήρωση των εντύπων."
    )

    # Φόρτωση μελών
    members_dict = load_members_dict()

    def member_selectbox(label: str, key: str) -> Optional[Dict]:
        """Selectbox επιλογής μέλους."""
        options = {0: "— Επιλογή μέλους —"}
        for mid, m in members_dict.items():
            options[mid] = f"{m.get('επώνυμο','')} {m.get('όνομα','')} ({m.get('τεκτονικός_βαθμός','')})"
        sel = st.selectbox(label, list(options.keys()),
                           format_func=lambda x: options[x], key=key)
        return members_dict.get(sel) if sel else None

    # Συλλογή δεδομένων ανά group
    collected: Dict[str, Any] = {}
    member_refs: Dict[str, Optional[Dict]] = {}  # field_id → member dict

    for group in proc["input_groups"]:
        st.markdown(f"#### {group['title']}")
        cols = st.columns(2)
        col_idx = 0

        for field_id, field_label, field_type, default in group["fields"]:
            if field_type == "auto":
                continue  # will be filled from member_refs

            with cols[col_idx % 2]:
                if field_type == "text":
                    collected[field_id] = st.text_input(
                        field_label, value=str(default), key=f"inp_{field_id}")
                elif field_type == "textarea":
                    collected[field_id] = st.text_area(
                        field_label, value=str(default), height=80, key=f"inp_{field_id}")
                elif field_type == "date":
                    try:
                        d = date.fromisoformat(default) if default else date.today()
                    except Exception:
                        d = date.today()
                    collected[field_id] = st.date_input(
                        field_label, value=d, key=f"inp_{field_id}"
                    ).strftime("%d/%m/%Y")
                elif field_type == "number":
                    collected[field_id] = st.number_input(
                        field_label, min_value=0, value=int(default) if default else 0,
                        key=f"inp_{field_id}")
                elif field_type == "select":
                    opts = default if isinstance(default, list) else [default]
                    collected[field_id] = st.selectbox(
                        field_label, opts, key=f"inp_{field_id}")
                elif field_type == "member":
                    sel_member = member_selectbox(field_label, key=f"inp_{field_id}")
                    member_refs[field_id] = sel_member
                    if sel_member:
                        collected[field_id] = member_display_name(sel_member)
                    else:
                        collected[field_id] = ""

            col_idx += 1

        # Auto-fill από member_refs
        for field_id, field_label, field_type, default in group["fields"]:
            if field_type != "auto":
                continue
            # Find the "member" field in same group
            base_id = field_id.rsplit("_", 1)[0]  # e.g. "yp_eponimo" → look for "yp_melos"
            # Try to find a member ref with matching prefix
            ref_member = None
            for mid_key, mref in member_refs.items():
                if mid_key.startswith(base_id.rsplit("_", 1)[0]) or base_id.startswith(mid_key.rsplit("_", 1)[0]):
                    ref_member = mref
                    break
                # Simple fallback: first member ref in group
                if not ref_member and mref:
                    ref_member = mref

            if ref_member:
                suffix = field_id.split("_")[-1]
                mapping = {
                    "eponimo":   ref_member.get("επώνυμο", ""),
                    "onoma":     ref_member.get("όνομα", ""),
                    "patronimo": ref_member.get("πατρώνυμο", ""),
                    "arith":     ref_member.get("αρ_μητρώου_στοάς", ""),
                    "arith_ms":  ref_member.get("αρ_μητρώου_μσ", ""),
                    "vathmος":   ref_member.get("τεκτονικός_βαθμός", ""),
                    "im_mithsis":ref_member.get("ημ_μύησης", ""),
                    "email":     ref_member.get("email", ""),
                }
                collected[field_id] = mapping.get(suffix, "")

        st.markdown("---")

    # Επιπλέον σχόλια
    extra_notes = st.text_area(
        "📝 Επιπλέον σχόλια / οδηγίες για τη συμπλήρωση",
        height=70,
        placeholder="π.χ. ειδικές παρατηρήσεις, σημειώσεις για Claude…",
        key="extra_notes",
    )

    # ── ΚΟΥΜΠΙ ΣΥΜΠΛΗΡΩΣΗΣ ───────────────────────────────────
    st.markdown("---")

    if st.button(
        "🤖 Αυτόματη Συμπλήρωση από Μητρώο",
        type="primary",
        use_container_width=True,
        key="run_fill",
    ):
        # Αποθηκεύουμε τα fill specs στο session_state
        stoaa_data = {"name": STOAA_NAME, "number": STOAA_NUMBER, "anatoli": STOAA_ANATOLI}

        data_for_fill = {}
        for k, v in collected.items():
            data_for_fill[k] = str(v) if v is not None else ""

        # Για εκλογές: φτιάχνουμε spec αξιωματικών
        if proc["code"] == "ekloges":
            officials_spec = []
            for axioma_text, field_key in [
                ("ΣΕΒΑΣΜΙΟΣ",            "ax_sev"),
                ("Α΄ ΕΠΟΠΤΗΣ",           "ax_a_ep"),
                ("Β΄ ΕΠΟΠΤΗΣ",           "ax_b_ep"),
                ("ΡΗΤΩΡ",                "ax_rhtor"),
                ("ΓΡΑΜΜ. - ΣΦΡΑΓΙΔ.",   "ax_gramm"),
                ("Α΄ ΔΟΚΙΜΑΣΤΗΣ",        "ax_a_dok"),
                ("ΤΑΜΙΑΣ",               "ax_tamias"),
                ("ΕΛΕΟΝΟΜΟΣ",            "ax_eleon"),
                ("ΤΕΛΕΤΑΡΧΗΣ",           "ax_tel"),
                ("ΣΤΕΓΑΣΤΗΣ",            "ax_steg"),
                ("Β΄ ΔΟΚΙΜΑΣΤΗΣ",        "ax_b_dok"),
                ("ΑΡΧΙΤ. - ΑΡΧΙΤΡ.",    "ax_arxitekt"),
                ("ΑΡΧΕΙΟΦ. - ΒΙΒΛΙΟΦ.", "ax_arxif"),
                ("ΞΙΦΟΦ. - ΣΗΜΑΙΟΦ.",   "ax_xifok"),
                ("ΡΗΤΩΡ (Πρόσθ.)",       "pr_rhtor"),
                ("ΓΡΑΜΜ. - ΣΦΡΑΓΙΔ. (Πρόσθ.)", "pr_gramm"),
                ("ΤΑΜΙΑΣ (Πρόσθ.)",      "pr_tamias"),
                ("ΕΛΕΟΝΟΜΟΣ (Πρόσθ.)",   "pr_eleon"),
                ("ΤΕΛΕΤΑΡΧΗΣ (Πρόσθ.)",  "pr_tel"),
                ("ΜΕΛΟΣ ΕΞ.ΕΠ. Α",       "ex_1"),
                ("ΜΕΛΟΣ ΕΞ.ΕΠ. Β",       "ex_2"),
                ("ΜΕΛΟΣ ΕΞ.ΕΠ. Γ",       "ex_3"),
            ]:
                m = member_refs.get(field_key)
                officials_spec.append({
                    "Αξίωμα": axioma_text,
                    "Επώνυμο": m.get("επώνυμο", "") if m else "",
                    "Όνομα": m.get("όνομα", "") if m else "",
                    "Πατρώνυμο": m.get("πατρώνυμο", "") if m else "",
                    "Α.Μ.Μ.Σ.": str(m.get("αρ_μητρώου_μσ", "") or "") if m else "",
                    "Κινητό": m.get("κινητό", "") or m.get("κινητο", "") if m else "",
                    "Email": m.get("email", "") if m else "",
                    "Διεύθυνση": m.get("διεύθυνση", "") if m else "",
                    "Πόλη": m.get("πόλη", "") if m else "",
                    "_field_key": field_key,
                })

            st.session_state[f"fill_officials_{proc['code']}"] = officials_spec
            st.session_state[f"fill_header_{proc['code']}"] = {
                "Τεκτ. Διετία": data_for_fill.get("ekl_diettia", "2024-2026"),
                "Ημ/νία Αρχαιρεσιών": data_for_fill.get("ekl_imerominia", ""),
                "Ψηφίσαντες": data_for_fill.get("ekl_psifisantes", ""),
                "Παρόντες": data_for_fill.get("ekl_paronton", ""),
            }

        st.session_state[f"fill_data_{proc['code']}"] = data_for_fill
        st.session_state[f"fill_members_{proc['code']}"] = {
            k: v for k, v in member_refs.items() if v
        }
        st.session_state[f"fill_stoaa_{proc['code']}"] = stoaa_data
        st.success("✅ Δεδομένα φορτώθηκαν. Ελέγξτε/διορθώστε παρακάτω και κατεβάστε.")

    # ══════════════════════════════════════════════════════════
    # EDITOR — Εμφανίζεται αν υπάρχουν δεδομένα στο session_state
    # ══════════════════════════════════════════════════════════
    if f"fill_data_{proc['code']}" in st.session_state:
        st.markdown("---")
        st.markdown("### ✏️ Επεξεργασία & Λήψη Εντύπων")
        st.caption("Διορθώστε χειροκίνητα οποιοδήποτε πεδίο και μετά κατεβάστε.")

        stoaa_data = st.session_state.get(f"fill_stoaa_{proc['code']}",
                                          {"name": STOAA_NAME, "number": STOAA_NUMBER,
                                           "anatoli": STOAA_ANATOLI})
        data_for_fill = st.session_state[f"fill_data_{proc['code']}"]
        saved_members = st.session_state.get(f"fill_members_{proc['code']}", {})

        # ── ΕΚΛΟΓΕΣ: Editable grid αξιωματικών ──────────────
        if proc["code"] == "ekloges":
            header_data = st.session_state.get(f"fill_header_{proc['code']}", {})
            officials_spec = st.session_state.get(f"fill_officials_{proc['code']}", [])

            # Επεξεργασία header fields
            st.markdown("#### 🏛️ Στοιχεία Αρχαιρεσιών")
            hcols = st.columns(4)
            edited_header = {}
            for i, (k, v) in enumerate(header_data.items()):
                edited_header[k] = hcols[i % 4].text_input(k, value=str(v),
                                                             key=f"hdr_{proc['code']}_{k}")
            st.session_state[f"fill_header_{proc['code']}"] = edited_header

            # Επεξεργασία πίνακα αξιωματικών
            st.markdown("#### 👥 Πίνακας Αξιωματικών")
            st.caption("Επεξεργαστείτε απευθείας οποιοδήποτε κελί.")

            import pandas as pd
            df = pd.DataFrame([
                {k: v for k, v in row.items() if k != "_field_key"}
                for row in officials_spec
            ])

            edited_df = st.data_editor(
                df,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Αξίωμα": st.column_config.TextColumn("Αξίωμα", disabled=True, width="medium"),
                    "Επώνυμο": st.column_config.TextColumn("Επώνυμο", width="medium"),
                    "Όνομα": st.column_config.TextColumn("Όνομα", width="medium"),
                    "Πατρώνυμο": st.column_config.TextColumn("Πατρώνυμο", width="small"),
                    "Α.Μ.Μ.Σ.": st.column_config.TextColumn("Α.Μ.Μ.Σ.", width="small"),
                    "Κινητό": st.column_config.TextColumn("Κινητό", width="small"),
                    "Email": st.column_config.TextColumn("Email", width="medium"),
                    "Διεύθυνση": st.column_config.TextColumn("Διεύθυνση", width="medium"),
                    "Πόλη": st.column_config.TextColumn("Πόλη", width="small"),
                },
                key=f"editor_grid_{proc['code']}",
                num_rows="fixed",
            )

            # Αποθήκευση επεξεργασμένου πίνακα
            updated_officials = []
            for i, row in edited_df.iterrows():
                entry = row.to_dict()
                entry["_field_key"] = officials_spec[i]["_field_key"] if i < len(officials_spec) else ""
                updated_officials.append(entry)
            st.session_state[f"fill_officials_{proc['code']}"] = updated_officials

        else:
            # Για μη-εκλογικά: απλά text inputs
            st.markdown("#### 📝 Στοιχεία Εντύπου")
            edit_cols = st.columns(2)
            edited_data = {}
            for i, (k, v) in enumerate(data_for_fill.items()):
                if k == "extra_notes":
                    continue
                edited_data[k] = edit_cols[i % 2].text_input(
                    k, value=str(v), key=f"edit_{proc['code']}_{k}")
            st.session_state[f"fill_data_{proc['code']}"] = edited_data

        # ── ΚΟΥΜΠΙ ΔΗΜΙΟΥΡΓΙΑΣ ΕΝΤΥΠΩΝ ───────────────────────
        st.markdown("---")
        st.markdown("### 📄 Δημιουργία & Λήψη Εντύπων")

        if st.button("📄 Δημιουργία εντύπων από τα παραπάνω δεδομένα",
                     type="primary", use_container_width=True, key="gen_from_edit"):

            filled_files: List[Tuple[str, bytes]] = []
            progress = st.progress(0)
            n_forms = len(proc["forms"])

            # Ανακατασκευή member_refs από τον edited grid (για εκλογές)
            if proc["code"] == "ekloges":
                updated_spec = st.session_state.get(f"fill_officials_{proc['code']}", [])
                # Φτιάχνουμε ψεύτικα member dicts από τον edited πίνακα
                rebuilt_members: Dict[str, Optional[Dict]] = {}
                for row in updated_spec:
                    fk = row.get("_field_key", "")
                    if fk:
                        rebuilt_members[fk] = {
                            "επώνυμο":       row.get("Επώνυμο", ""),
                            "όνομα":         row.get("Όνομα", ""),
                            "πατρώνυμο":     row.get("Πατρώνυμο", ""),
                            "αρ_μητρώου_μσ": row.get("Α.Μ.Μ.Σ.", ""),
                            "κινητό":        row.get("Κινητό", ""),
                            "email":         row.get("Email", ""),
                            "διεύθυνση":     row.get("Διεύθυνση", ""),
                            "πόλη":          row.get("Πόλη", ""),
                        }
                # Ενημέρωση header data
                hdr = st.session_state.get(f"fill_header_{proc['code']}", {})
                final_collected = dict(data_for_fill)
                final_collected["ekl_diettia"] = hdr.get("Τεκτ. Διετία", "2024-2026")
                final_collected["ekl_imerominia"] = hdr.get("Ημ/νία Αρχαιρεσιών", "")
                final_collected["ekl_psifisantes"] = hdr.get("Ψηφίσαντες", "")
                final_collected["ekl_paronton"] = hdr.get("Παρόντες", "")
            else:
                rebuilt_members = saved_members
                final_collected = st.session_state.get(f"fill_data_{proc['code']}", data_for_fill)

            for i, form_info in enumerate(proc["forms"]):
                path = os.path.join(FORMS_ROOT, form_info["file"])
                progress.progress(int((i / n_forms) * 85))

                if not os.path.exists(path):
                    st.warning(f"⚠️ Δεν βρέθηκε: {path}")
                    continue

                try:
                    with st.spinner(f"Συμπλήρωση «{form_info['name']}»…"):
                        filled_bytes = fill_docx_smart(
                            template_path=path,
                            form_file_key=form_info["file"],
                            collected=final_collected,
                            members_by_field=rebuilt_members,
                            stoaa_data=stoaa_data,
                        )

                    fname = os.path.basename(form_info["file"])
                    filled_files.append((fname, filled_bytes))

                    st.download_button(
                        f"⬇️ {form_info['name']}",
                        data=filled_bytes,
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_filled_{i}_{proc['code']}",
                        use_container_width=True,
                    )
                    st.success(f"✅ {form_info['name']}")

                except Exception as e:
                    st.error(f"❌ {form_info['name']}: {e}")

            progress.progress(100)

            if len(filled_files) > 1:
                zip_data = create_zip(filled_files)
                st.download_button(
                    "📦 Λήψη όλων (.zip)",
                    data=zip_data,
                    file_name=f"εντυπα_{proc['code']}_{date.today().strftime('%Y%m%d')}.zip",
                    mime="application/zip",
                    use_container_width=True,
                    key=f"dl_zip_{proc['code']}",
                )

    # ── ΟΔΗΓΙΕΣ ΑΠΟΣΤΟΛΗΣ ─────────────────────────────────────
    with st.expander("📬 Οδηγίες Αποστολής στη Μεγ. Γεν. Γραμματεία"):
        st.markdown("""
**Αποστολή εντύπων:**
- **Πρωτότυπα**: αυτοπροσώπως (Σεβ. ή Γραμμ.) ή courier ή ταχυδρομείο
- **Αντίγραφα**: e-mail μέσω **επίσημης** ηλεκτρονικής διεύθυνσης Στοάς
- **Διεύθυνση**: Αχαρνών 19, 10438 Αθήναι
- **Email**: megaligrammatia.gl@grandlodge.gr
- **Τηλ.**: 210 8229950

⚠️ **Χωρίς πρωτότυπα δεν εκδίδεται έγκριση από τη Μεγ. Στοά.**
        """)
